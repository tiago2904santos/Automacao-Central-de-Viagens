from __future__ import annotations

import json
from datetime import date, time
from io import BytesIO

from django.test import TestCase
from django.urls import reverse
from docx import Document

from viagens.documents.plano_trabalho import build_plano_trabalho_docx_bytes
from viagens.models import (
    Cidade,
    CoordenadorMunicipal,
    Estado,
    Oficio,
    PlanoTrabalho,
    PlanoTrabalhoAtividade,
    PlanoTrabalhoLocalAtuacao,
    PlanoTrabalhoMeta,
    PlanoTrabalhoRecurso,
    Trecho,
    Viajante,
)
from viagens.services.plano_trabalho import build_coordenacao_formatada


class PlanoTrabalhoEnxutoTests(TestCase):
    def setUp(self) -> None:
        self.estado_pr = Estado.objects.create(sigla="PR", nome="Parana")
        self.cidade_sede = Cidade.objects.create(nome="Curitiba", estado=self.estado_pr)
        self.cidade_destino = Cidade.objects.create(nome="Foz do Iguacu", estado=self.estado_pr)
        self.coordenador_admin = Viajante.objects.create(
            nome="Coordenador Administrativo",
            rg="12345678X",
            cpf="00000000000",
            cargo="Delegado",
        )
        self.servidor = Viajante.objects.create(
            nome="Servidor Operacional",
            rg="98765432X",
            cpf="11111111111",
            cargo="Agente",
        )

        self.oficio = Oficio.objects.create(
            oficio="501/2026",
            protocolo="121234567",
            assunto="Plano para atendimento itinerante",
            placa="ABC1234",
            modelo="Viatura",
            combustivel="Gasolina",
            motorista=self.servidor.nome,
            tipo_destino="INTERIOR",
            motivo="Atendimento ao publico",
            estado_sede=self.estado_pr,
            cidade_sede=self.cidade_sede,
            estado_destino=self.estado_pr,
            cidade_destino=self.cidade_destino,
            retorno_saida_data=date(2026, 3, 12),
            retorno_saida_hora=time(8, 0),
            retorno_chegada_data=date(2026, 3, 14),
            retorno_chegada_hora=time(18, 0),
        )
        self.oficio.viajantes.add(self.servidor)
        Trecho.objects.create(
            oficio=self.oficio,
            ordem=1,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_sede,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino,
            saida_data=date(2026, 3, 12),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 3, 12),
            chegada_hora=time(12, 0),
        )

    def _doc_text(self, payload: bytes) -> str:
        doc = Document(BytesIO(payload))
        chunks = [paragraph.text for paragraph in doc.paragraphs]
        for section in doc.sections:
            chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
            chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(chunks)

    def _base_payload(self) -> dict[str, str]:
        return {
            "numero": "",
            "ano": "2026",
            "sigla_unidade": "ASCOM",
            "programa_projeto": "PCPR na Comunidade",
            "destino": "Foz do Iguacu/PR",
            "solicitante": "Prefeitura Municipal",
            "contexto_solicitacao": "Atendimento solicitado para emissao de documentos.",
            "data_inicio": "2026-03-12",
            "data_fim": "2026-03-14",
            "horario_atendimento": "das 09h as 17h",
            "efetivo_formatado": "12 servidores.",
            "estrutura_apoio": "",
            "quantidade_servidores": "12",
            "composicao_diarias": "4 x 100% + 1 x 30%",
            "valor_unitario": "150,00",
            "valor_total_calculado": "",
            "coordenador_plano": str(self.coordenador_admin.id),
            "coordenador_municipal": "",
            "possui_coordenador_municipal": "nao",
            "coordenador_municipal_nome": "",
            "coordenador_municipal_cargo": "",
            "coordenador_municipal_cidade": "",
            "texto_override": "",
            "atividades_selecionadas": [
                "Confecção da Carteira de Identidade Nacional (CIN)"
            ],
            "metas_json": json.dumps(
                [
                    {
                        "descricao": "ampliar o acesso ao documento oficial de identificação civil, garantindo cidadania e inclusão social à população atendida."
                    }
                ]
            ),
            "atividades_json": json.dumps(
                [{"descricao": "Confecção da Carteira de Identidade Nacional (CIN)"}]
            ),
            "recursos_json": json.dumps([{"descricao": "Unidade movel da PCPR."}]),
            "locais_json": json.dumps(
                [{"data": "2026-03-12", "local": "Foz do Iguacu/PR"}]
            ),
        }

    def _create_plano_base(self) -> PlanoTrabalho:
        plano = PlanoTrabalho.objects.create(
            oficio=self.oficio,
            numero=5,
            ano=2026,
            sigla_unidade="ASCOM",
            programa_projeto="PCPR na Comunidade",
            destino="Foz do Iguacu/PR",
            solicitante="Prefeitura Municipal",
            contexto_solicitacao="Solicitacao formal do municipio.",
            local="Foz do Iguacu/PR",
            data_inicio=date(2026, 3, 12),
            data_fim=date(2026, 3, 14),
            horario_atendimento="das 09h as 17h",
            efetivo_formatado="12 servidores.",
            efetivo_por_dia=12,
            quantidade_servidores=12,
            composicao_diarias="4 x 100% + 1 x 30%",
            valor_unitario="150.00",
            valor_total_calculado="7740.00",
            coordenador_plano=self.coordenador_admin,
            coordenador_nome=self.coordenador_admin.nome,
            coordenador_cargo=self.coordenador_admin.cargo,
        )
        PlanoTrabalhoMeta.objects.create(
            plano=plano,
            ordem=1,
            descricao="ampliar o acesso ao documento oficial de identificação civil, garantindo cidadania e inclusão social à população atendida.",
        )
        PlanoTrabalhoAtividade.objects.create(
            plano=plano,
            ordem=1,
            descricao="Confecção da Carteira de Identidade Nacional (CIN)",
        )
        PlanoTrabalhoRecurso.objects.create(
            plano=plano,
            ordem=1,
            descricao="Unidade movel da PCPR.",
        )
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=1,
            data=date(2026, 3, 12),
            local="Foz do Iguacu/PR",
        )
        return plano

    def test_cenario_1_sem_coordenador_municipal(self) -> None:
        plano = self._create_plano_base()
        texto = build_coordenacao_formatada(plano)
        self.assertIn("Coordenadora Administrativa do Plano", texto)
        self.assertNotIn("Coordenador(a) Municipal", texto)

    def test_cenario_2_multiplos_locais_com_estrutura(self) -> None:
        plano = self._create_plano_base()
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=2,
            data=date(2026, 3, 13),
            local="Escola Estadual XYZ",
        )
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=3,
            data=date(2026, 3, 14),
            local="Distrito ABC",
        )
        plano.estrutura_apoio = "Unidade movel da PCPR equipada para atendimento."
        plano.save(update_fields=["estrutura_apoio", "updated_at"])

        text = self._doc_text(build_plano_trabalho_docx_bytes(self.oficio).getvalue())
        self.assertIn("Foz do Iguacu/PR, Escola Estadual XYZ e Distrito ABC", text)
        self.assertIn("Unidade movel da PCPR equipada para atendimento.", text)

    def test_cenario_3_com_coordenador_municipal(self) -> None:
        plano = self._create_plano_base()
        coord_municipal = CoordenadorMunicipal.objects.create(
            nome="Joana Silva",
            cargo="Secretaria Municipal",
            cidade="Foz do Iguacu",
            ativo=True,
        )
        plano.possui_coordenador_municipal = True
        plano.coordenador_municipal = coord_municipal
        plano.save(update_fields=["possui_coordenador_municipal", "coordenador_municipal", "updated_at"])

        text = build_coordenacao_formatada(plano)
        self.assertIn("Coordenadora Administrativa do Plano", text)
        self.assertIn("Coordenador(a) Municipal do Evento", text)
        self.assertIn("Joana Silva", text)

    def test_cenario_4_cadastro_inline_coordenador_municipal(self) -> None:
        payload = self._base_payload()
        payload.update(
            {
                "possui_coordenador_municipal": "sim",
                "coordenador_municipal_nome": "Carlos Pereira",
                "coordenador_municipal_cargo": "Diretor Municipal",
                "coordenador_municipal_cidade": "Foz do Iguacu",
            }
        )
        response = self.client.post(
            reverse("plano_trabalho_editar", args=[self.oficio.id]),
            payload,
        )
        if response.status_code != 302:
            form = response.context["form"]
            self.fail(form.errors.as_json())
        self.assertEqual(response.status_code, 302)

        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        self.assertTrue(plano.possui_coordenador_municipal)
        self.assertIsNotNone(plano.coordenador_municipal)
        self.assertEqual(plano.coordenador_municipal.nome, "Carlos Pereira")
        self.assertTrue(
            CoordenadorMunicipal.objects.filter(
                nome="Carlos Pereira",
                cidade="Foz do Iguacu",
            ).exists()
        )

    def test_cenario_5_erro_quando_coordenador_municipal_ausente(self) -> None:
        payload = self._base_payload()
        payload.update({"possui_coordenador_municipal": "sim"})
        response = self.client.post(
            reverse("plano_trabalho_editar", args=[self.oficio.id]),
            payload,
        )
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("coordenador_municipal", form.errors)

    def test_cenario_6_template_final_sem_placeholders(self) -> None:
        self._create_plano_base()
        text = self._doc_text(build_plano_trabalho_docx_bytes(self.oficio).getvalue())
        self.assertIn("PLANO DE TRABALHO", text)
        self.assertIn("BREVE CONTEXTUALIZAÇÃO", text)
        self.assertIn("METAS ESTABELECIDAS", text)
        self.assertIn("ATIVIDADES A SEREM DESENVOLVIDAS", text)
        self.assertIn("VALOR TOTAL DO PLANO", text)
        self.assertIn("COORDENADOR DO EVENTO", text)
        self.assertNotIn("{{", text)
        self.assertNotIn("}}", text)
