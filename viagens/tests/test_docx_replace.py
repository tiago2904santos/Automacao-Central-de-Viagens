import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from django.test import SimpleTestCase, override_settings

from viagens.documents.document import (
    _diagnose_docx_xml_on_open_error,
    _sub_placeholders,
    replace_placeholders_in_paragraph,
)


class DocxReplaceTests(SimpleTestCase):
    def test_replace_preserves_run_formatting(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("RG: {{rg}}")
        run.bold = True
        run.font.size = Pt(10.5)

        replace_placeholders_in_paragraph(paragraph, {"rg": "123"})

        self.assertEqual(paragraph.runs[0].text, "RG: 123")
        self.assertTrue(paragraph.runs[0].bold)
        self.assertEqual(paragraph.runs[0].font.size, Pt(10.5))

    def test_replace_across_runs_keeps_styles(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        first = paragraph.add_run("CPF: {{")
        first.bold = True
        second = paragraph.add_run("cpf}}")
        second.bold = True

        replace_placeholders_in_paragraph(paragraph, {"cpf": "999"})

        self.assertEqual(paragraph.runs[0].text, "CPF: 999")
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(paragraph.runs[1].bold)

    def test_replace_accepts_placeholder_with_spaced_closing_brace(self):
        text = "Rodape: {{unidade_rodape} }"
        replaced = _sub_placeholders(text, {"unidade_rodape": "DPC"})
        self.assertEqual(replaced, "Rodape: DPC")

    def test_replace_sanitizes_invalid_xml_control_chars(self):
        text = "Nome: {{nome}}"
        replaced = _sub_placeholders(text, {"nome": "A\x00B\x01C\x0bD\x0cE\tF\nG\rH&<>"})
        self.assertEqual(replaced, "Nome: ABCDE\tF\nG\rH&<>")

    def test_diagnose_docx_xml_reports_broken_document_xml(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            base_dir = Path(tmpdir)
            docx_path = base_dir / "broken.docx"
            with zipfile.ZipFile(docx_path, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("word/document.xml", "<root><a></root>")
                archive.writestr("word/styles.xml", "<root/>")
                archive.writestr("word/numbering.xml", "<root/>")
                archive.writestr("word/header1.xml", "<root/>")
                archive.writestr("word/footer1.xml", "<root/>")

            with override_settings(BASE_DIR=str(base_dir), DEBUG=True):
                unzip_dir, failures = _diagnose_docx_xml_on_open_error(
                    docx_path,
                    oficio_id=38,
                )

            self.assertIsNotNone(unzip_dir)
            self.assertTrue(Path(unzip_dir).exists())
            self.assertTrue(any("word/document.xml" in failure for failure in failures))
