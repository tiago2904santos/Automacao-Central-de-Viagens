from __future__ import annotations

from datetime import date
from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from viagens.models import Oficio, OrdemServico, Trecho, get_next_ordem_num
from viagens.services.oficio_config import get_oficio_config

MESES_PTBR = {
    1: "janeiro",
    2: "fevereiro",
    3: "marco",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}

HEADER_LINES = (
    "SECRETARIA DE ESTADO DA SEGURANÇA PÚBLICA",
    "POLÍCIA CIVIL DO PARANÁ",
    "ASSESSORIA DE COMUNICAÇÃO SOCIAL",
)


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _format_data_extenso(value: date | None) -> str:
    if not value:
        return ""
    mes = MESES_PTBR.get(value.month, str(value.month))
    return f"{value.day} de {mes} de {value.year}"


def _resolve_periodo(oficio: Oficio, trechos: list[Trecho]) -> tuple[date, date]:
    datas_inicio = [trecho.saida_data for trecho in trechos if trecho.saida_data]
    datas_fim = [
        trecho.chegada_data or trecho.saida_data
        for trecho in trechos
        if trecho.chegada_data or trecho.saida_data
    ]
    if oficio.retorno_saida_data:
        datas_fim.append(oficio.retorno_saida_data)
    if oficio.retorno_chegada_data:
        datas_fim.append(oficio.retorno_chegada_data)

    hoje = timezone.localdate()
    data_inicio = min(datas_inicio) if datas_inicio else hoje
    data_fim = max(datas_fim) if datas_fim else data_inicio
    if data_fim < data_inicio:
        data_fim = data_inicio
    return data_inicio, data_fim


def _resolve_destinos(trechos: list[Trecho], oficio: Oficio) -> str:
    destinos: list[str] = []
    seen: set[str] = set()
    for trecho in trechos:
        if trecho.destino_cidade and trecho.destino_estado:
            label = f"{trecho.destino_cidade.nome}/{trecho.destino_estado.sigla}"
        elif trecho.destino_cidade:
            label = trecho.destino_cidade.nome
        else:
            label = ""
        if label and label not in seen:
            seen.add(label)
            destinos.append(label)
    if not destinos:
        if oficio.cidade_destino and oficio.estado_destino:
            return f"{oficio.cidade_destino.nome}/{oficio.estado_destino.sigla}"
        if oficio.cidade_destino:
            return oficio.cidade_destino.nome
        return "Curitiba/PR"
    if len(destinos) == 1:
        return destinos[0]
    if len(destinos) == 2:
        return f"{destinos[0]} e {destinos[1]}"
    return f"{', '.join(destinos[:-1])} e {destinos[-1]}"


def _resolve_viajantes(oficio: Oficio) -> str:
    itens = [
        f"{viajante.nome} ({viajante.cargo})"
        for viajante in oficio.viajantes.all().order_by("nome")
        if (viajante.nome or "").strip()
    ]
    if not itens:
        return "servidores designados"
    if len(itens) == 1:
        return itens[0]
    if len(itens) == 2:
        return f"{itens[0]} e {itens[1]}"
    return f"{', '.join(itens[:-1])} e {itens[-1]}"


def _ensure_ordem_servico(oficio: Oficio) -> OrdemServico:
    try:
        return oficio.ordem_servico
    except OrdemServico.DoesNotExist:
        pass

    cfg = get_oficio_config()
    ano = int(oficio.ano or timezone.localdate().year)
    determinante_nome = ""
    determinante_cargo = ""
    if getattr(cfg, "assinante", None):
        determinante_nome = (cfg.assinante.nome or "").strip()
        determinante_cargo = (cfg.assinante.cargo or "").strip()

    return OrdemServico.objects.create(
        oficio=oficio,
        numero=get_next_ordem_num(ano),
        ano=ano,
        referencia="Diligências",
        determinante_nome=determinante_nome,
        determinante_cargo=determinante_cargo,
        finalidade=(oficio.motivo or "").strip(),
    )


def build_ordem_servico_docx_bytes(oficio: Oficio) -> BytesIO:
    trechos = list(oficio.trechos.select_related("destino_cidade", "destino_estado").order_by("ordem"))
    ordem = _ensure_ordem_servico(oficio)
    cfg = get_oficio_config()

    determinante_nome = (ordem.determinante_nome or "").strip()
    determinante_cargo = (ordem.determinante_cargo or "").strip()
    if (not determinante_nome or not determinante_cargo) and getattr(cfg, "assinante", None):
        if not determinante_nome:
            determinante_nome = (cfg.assinante.nome or "").strip()
        if not determinante_cargo:
            determinante_cargo = (cfg.assinante.cargo or "").strip()

    data_inicio, data_fim = _resolve_periodo(oficio, trechos)
    destinos = _resolve_destinos(trechos, oficio)
    viajantes_text = _resolve_viajantes(oficio)

    finalidade = (ordem.finalidade or "").strip()
    if not finalidade:
        finalidade = "para atendimento das atividades institucionais."
    elif not finalidade.endswith("."):
        finalidade = f"{finalidade}."

    if ordem.texto_override.strip():
        corpo = ordem.texto_override.strip()
    else:
        corpo = (
            f"Eu, {determinante_nome or '---'}, {determinante_cargo or '---'}, "
            f"no uso de minhas atribuições, DETERMINO:\n\n"
            f"O deslocamento de {viajantes_text} para o município de {destinos}, "
            f"no período de {_format_data_extenso(data_inicio)} a {_format_data_extenso(data_fim)}, "
            f"{finalidade}"
        )

    doc = Document()
    _set_default_style(doc)

    for line in HEADER_LINES:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].bold = True

    titulo = doc.add_paragraph(f"ORDEM DE SERVIÇO Nº {ordem.numero}/{ordem.ano}/ASCOM")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if titulo.runs:
        titulo.runs[0].bold = True

    ref = doc.add_paragraph(f"Ref.: {(ordem.referencia or 'Diligências').strip()}")
    ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if ref.runs:
        ref.runs[0].bold = True

    doc.add_paragraph("")
    for bloco in corpo.split("\n"):
        p = doc.add_paragraph(bloco)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("")
    hoje = timezone.localdate()
    data_assinatura = doc.add_paragraph(
        f"Curitiba, {hoje.day} de {MESES_PTBR.get(hoje.month, hoje.month)} de {hoje.year}."
    )
    data_assinatura.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")
    assinatura_nome = doc.add_paragraph(determinante_nome or "\u2014")
    assinatura_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if assinatura_nome.runs:
        assinatura_nome.runs[0].bold = True
    assinatura_cargo = doc.add_paragraph(determinante_cargo or "\u2014")
    assinatura_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
