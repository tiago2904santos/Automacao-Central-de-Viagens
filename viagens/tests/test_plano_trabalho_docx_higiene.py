from __future__ import annotations

from django.test import TestCase
from docx import Document

from viagens.documents.document import remove_optional_placeholder_paragraphs
from viagens.services.plano_trabalho import (
    apply_text_hygiene,
    resolve_plano_docx_placeholders,
)


class PlanoTrabalhoDocxHigieneTests(TestCase):
    def test_resolve_aliases_legados_do_template(self) -> None:
        template_keys = [
            "coordenação formatada",
            "metas_formatada",
            "atividades_formatada",
            "valor_total_por_extenso",
        ]
        placeholders = {
            "coordenacao_formatada": "Coordenacao do plano",
            "metas_formatadas": "• Meta 1",
            "atividades_formatada": "• Atividade 1",
            "valor_total_extenso": "cem reais",
        }
        resolved, missing_required = resolve_plano_docx_placeholders(
            template_keys,
            placeholders,
        )
        self.assertEqual(missing_required, [])
        self.assertEqual(resolved["coordenação formatada"], "Coordenacao do plano")
        self.assertEqual(resolved["metas_formatada"], "• Meta 1")
        self.assertEqual(resolved["atividades_formatada"], "• Atividade 1")
        self.assertEqual(resolved["valor_total_por_extenso"], "cem reais")

    def test_apply_text_hygiene_remove_ruido_das_das(self) -> None:
        value = apply_text_hygiene("das das 17h  às   02h")
        self.assertEqual(value, "das 17h às 02h")

    def test_apply_text_hygiene_normaliza_nao_informado(self) -> None:
        self.assertEqual(apply_text_hygiene("Nao informado."), "Não informado.")

    def test_remove_paragrafo_placeholder_opcional_vazio(self) -> None:
        doc = Document()
        doc.add_paragraph("{{unidade_movel}}")
        doc.add_paragraph("Linha fixa")
        remove_optional_placeholder_paragraphs(
            doc,
            resolved_mapping={"unidade_movel": ""},
            optional_placeholders={"unidade_movel"},
        )
        texts = [p.text for p in doc.paragraphs]
        self.assertEqual(texts, ["Linha fixa"])
