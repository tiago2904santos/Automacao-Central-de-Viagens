from __future__ import annotations

import json
from datetime import date, time
from io import BytesIO

from django.test import TestCase
from django.urls import reverse
from docx import Document

from viagens.documents.plano_trabalho import build_plano_trabalho_docx_bytes
from viagens.models import (
    Cidade,
    Estado,
    Oficio,
    OficioConfig,
    PlanoTrabalho,
    PlanoTrabalhoAtividade,
    PlanoTrabalhoLocalAtuacao,
    PlanoTrabalhoMeta,
    PlanoTrabalhoRecurso,
    Trecho,
    Viajante,
)
from viagens.services.plano_trabalho import (
    ATIVIDADES_ORDEM_FIXA,
    META_POR_ATIVIDADE,
    build_plano_placeholders,
    format_lista_portugues,
    format_periodo_evento_extenso,
    metas_from_atividades,
)


class PlanoTrabalhoPaginas01e02Tests(TestCase):
    def setUp(self) -> None:
        self.estado_pr = Estado.objects.create(sigla="PR", nome="Parana")
        self.cidade_sede = Cidade.objects.create(nome="Curitiba", estado=self.estado_pr)
        self.cidade_destino = Cidade.objects.create(nome="Ponta Grossa", estado=self.estado_pr)
        self.coordenador = Viajante.objects.create(
            nome="Ana Coordenadora",
            rg="12345678X",
            cpf="00000000000",
            cargo="Delegada",
        )
        self.servidor = Viajante.objects.create(
            nome="Servidor Teste",
            rg="99999999X",
            cpf="11111111111",
            cargo="Agente",
        )
        self.oficio = Oficio.objects.create(
            oficio="701/2026",
            protocolo="121234567",
            assunto="Plano festas e exposicao",
            placa="ABC1234",
            modelo="Viatura",
            combustivel="Gasolina",
            motorista=self.servidor.nome,
            tipo_destino="INTERIOR",
            motivo="Atendimento",
            estado_sede=self.estado_pr,
            cidade_sede=self.cidade_sede,
            estado_destino=self.estado_pr,
            cidade_destino=self.cidade_destino,
        )
        self.oficio.viajantes.add(self.servidor)
        Trecho.objects.create(
            oficio=self.oficio,
            ordem=1,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_sede,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino,
            saida_data=date(2026, 2, 13),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 2, 13),
            chegada_hora=time(12, 0),
        )

    def _base_payload(self) -> dict[str, object]:
        atividade = ATIVIDADES_ORDEM_FIXA[0]
        return {
            "numero": "",
            "ano": "2026",
            "sigla_unidade": "ASCOM",
            "programa_projeto": "PCPR na Comunidade",
            "destino": "Ponta Grossa/PR",
            "solicitante": "Joao da Silva",
            "contexto_solicitacao": "Solicitacao municipal",
            "data_inicio": "2026-02-13",
            "data_fim": "2026-02-16",
            "horario_atendimento": "09h às 17h",
            "efetivo_formatado": "5 servidores.",
            "estrutura_apoio": "",
            "quantidade_servidores": "5",
            "composicao_diarias": "4 x 100% + 1 x 30%",
            "valor_unitario": "150,00",
            "valor_total_calculado": "",
            "coordenador_plano": str(self.coordenador.id),
            "coordenador_municipal": "",
            "possui_coordenador_municipal": "nao",
            "coordenador_municipal_nome": "",
            "coordenador_municipal_cargo": "",
            "coordenador_municipal_cidade": "",
            "texto_override": "",
            "atividades_selecionadas": [atividade],
            "atividades_json": json.dumps([{"descricao": atividade}], ensure_ascii=False),
            "metas_json": json.dumps(
                [{"descricao": META_POR_ATIVIDADE[atividade]}],
                ensure_ascii=False,
            ),
            "recursos_json": json.dumps([{"descricao": "Unidade movel da PCPR."}]),
            "locais_json": json.dumps(
                [
                    {"data": "2026-02-13", "local": "Praça Central"},
                    {"data": "2026-02-14", "local": "Ginasio Municipal"},
                ],
                ensure_ascii=False,
            ),
        }

    def _build_cfg(self) -> OficioConfig:
        cfg = OficioConfig()
        cfg.unidade_nome = "ASSESSORIA DE COMUNICACAO SOCIAL"
        cfg.origem_nome = "POLICIA CIVIL DO PARANA"
        cfg.cidade = "Curitiba"
        cfg.uf = "PR"
        cfg.telefone = "(41) 3235-6476"
        cfg.email = "ascom@pc.pr.gov.br"
        cfg.assinante = self.coordenador
        return cfg

    def _doc_text(self, payload: bytes) -> str:
        doc = Document(BytesIO(payload))
        chunks = [paragraph.text for paragraph in doc.paragraphs]
        for section in doc.sections:
            chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
            chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(chunks)

    def test_pagina_01_dias_evento_extenso_um_dia(self) -> None:
        value = format_periodo_evento_extenso(date(2026, 2, 13), date(2026, 2, 13))
        self.assertEqual(value, "13 de fevereiro de 2026")

    def test_pagina_01_dias_evento_extenso_intervalo_mes_igual(self) -> None:
        value = format_periodo_evento_extenso(date(2026, 2, 13), date(2026, 2, 16))
        self.assertEqual(value, "de 13 a 16 de fevereiro de 2026")

    def test_pagina_01_dias_evento_extenso_intervalo_meses_diferentes(self) -> None:
        value = format_periodo_evento_extenso(date(2026, 3, 30), date(2026, 4, 2))
        self.assertEqual(value, "de 30 de março a 2 de abril de 2026")

    def test_pagina_01_dias_evento_extenso_intervalo_anos_diferentes(self) -> None:
        value = format_periodo_evento_extenso(date(2026, 12, 30), date(2027, 1, 2))
        self.assertEqual(value, "de 30 de dezembro de 2026 a 2 de janeiro de 2027")

    def test_pagina_01_locais_formatado_um_dois_tres(self) -> None:
        self.assertEqual(format_lista_portugues(["Praça Central"]), "Praça Central")
        self.assertEqual(
            format_lista_portugues(["Praça Central", "Ginasio Municipal"]),
            "Praça Central e Ginasio Municipal",
        )
        self.assertEqual(
            format_lista_portugues(["Praça Central", "Ginasio Municipal", "Centro Cultural"]),
            "Praça Central, Ginasio Municipal e Centro Cultural",
        )

    def test_pagina_01_placeholders_no_contexto_docx(self) -> None:
        atividade = ATIVIDADES_ORDEM_FIXA[0]
        plano = PlanoTrabalho.objects.create(
            oficio=self.oficio,
            numero=1,
            ano=2026,
            sigla_unidade="ASCOM",
            destino="Ponta Grossa/PR",
            solicitante="Joao da Silva",
            local="Praça Central",
            data_inicio=date(2026, 2, 13),
            data_fim=date(2026, 2, 16),
            horario_atendimento="09h às 17h",
            efetivo_formatado="5 servidores.",
            quantidade_servidores=5,
            composicao_diarias="4 x 100% + 1 x 30%",
            valor_unitario="150.00",
            valor_total_calculado="3225.00",
            coordenador_plano=self.coordenador,
        )
        PlanoTrabalhoAtividade.objects.create(plano=plano, ordem=1, descricao=atividade)
        PlanoTrabalhoMeta.objects.create(
            plano=plano,
            ordem=1,
            descricao=META_POR_ATIVIDADE[atividade],
        )
        PlanoTrabalhoRecurso.objects.create(plano=plano, ordem=1, descricao="Unidade movel da PCPR.")
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=1,
            data=date(2026, 2, 13),
            local="Praça Central",
        )
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=2,
            data=date(2026, 2, 14),
            local="Ginasio Municipal",
        )

        context = build_plano_placeholders(plano, self.oficio, self._build_cfg())
        self.assertEqual(context["dias_evento_extenso"], "de 13 a 16 de fevereiro de 2026")
        self.assertEqual(context["locais_formatado"], "Praça Central e Ginasio Municipal")
        self.assertEqual(context["horario_atendimento"], "das 09h às 17h")
        self.assertEqual(context["quantidade_de_servidores"], "5")
        self.assertEqual(context["solicitante"], "Joao da Silva")

    def test_pagina_02_uma_atividade_gera_uma_meta(self) -> None:
        atividade = [ATIVIDADES_ORDEM_FIXA[2]]
        metas = metas_from_atividades(atividade)
        self.assertEqual(len(metas), 1)
        self.assertEqual(metas[0], META_POR_ATIVIDADE[atividade[0]])

    def test_pagina_02_varias_atividades_sem_duplicidade_ordem_fixa(self) -> None:
        atividades = [
            ATIVIDADES_ORDEM_FIXA[7],
            ATIVIDADES_ORDEM_FIXA[0],
            ATIVIDADES_ORDEM_FIXA[7],
            ATIVIDADES_ORDEM_FIXA[4],
        ]
        metas = metas_from_atividades(atividades)
        self.assertEqual(
            metas,
            [
                META_POR_ATIVIDADE[ATIVIDADES_ORDEM_FIXA[0]],
                META_POR_ATIVIDADE[ATIVIDADES_ORDEM_FIXA[4]],
                META_POR_ATIVIDADE[ATIVIDADES_ORDEM_FIXA[7]],
            ],
        )

    def test_pagina_02_remocao_atividade_remove_meta(self) -> None:
        atividade_a = ATIVIDADES_ORDEM_FIXA[0]
        atividade_b = ATIVIDADES_ORDEM_FIXA[1]
        metas_completas = metas_from_atividades([atividade_a, atividade_b])
        metas_reduzidas = metas_from_atividades([atividade_b])
        self.assertEqual(len(metas_completas), 2)
        self.assertEqual(metas_reduzidas, [META_POR_ATIVIDADE[atividade_b]])

    def test_pagina_02_contexto_docx_contem_atividades_e_metas_formatadas(self) -> None:
        atividade_a = ATIVIDADES_ORDEM_FIXA[0]
        atividade_b = ATIVIDADES_ORDEM_FIXA[6]
        plano = PlanoTrabalho.objects.create(
            oficio=self.oficio,
            numero=2,
            ano=2026,
            sigla_unidade="ASCOM",
            destino="Ponta Grossa/PR",
            solicitante="Joao da Silva",
            local="Praça Central",
            data_inicio=date(2026, 2, 13),
            data_fim=date(2026, 2, 16),
            horario_atendimento="das 09h às 17h",
            efetivo_formatado="5 servidores.",
            quantidade_servidores=5,
            composicao_diarias="4 x 100% + 1 x 30%",
            valor_unitario="150.00",
            valor_total_calculado="3225.00",
            coordenador_plano=self.coordenador,
        )
        PlanoTrabalhoAtividade.objects.create(plano=plano, ordem=1, descricao=atividade_b)
        PlanoTrabalhoAtividade.objects.create(plano=plano, ordem=2, descricao=atividade_a)
        PlanoTrabalhoRecurso.objects.create(plano=plano, ordem=1, descricao="Unidade movel da PCPR.")
        PlanoTrabalhoLocalAtuacao.objects.create(plano=plano, ordem=1, local="Praça Central")

        context = build_plano_placeholders(plano, self.oficio, self._build_cfg())
        self.assertIn("• Confecção da Carteira de Identidade Nacional (CIN)", context["atividades_formatada"])
        self.assertIn("• Exposição de material tático", context["atividades_formatada"])
        self.assertIn("• ampliar o acesso ao documento oficial de identificação civil", context["metas_formatadas"])
        self.assertIn("• apresentar equipamentos utilizados nas atividades policiais", context["metas_formatadas"])

    def test_pagina_02_validacao_bloqueia_sem_atividade(self) -> None:
        payload = self._base_payload()
        payload["atividades_selecionadas"] = []
        payload["atividades_json"] = "[]"
        payload["metas_json"] = "[]"
        response = self.client.post(reverse("plano_trabalho_editar", args=[self.oficio.id]), payload)
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("atividades_selecionadas", form.errors)

    def test_pagina_02_persistencia_e_reabertura_marcas(self) -> None:
        atividade_a = ATIVIDADES_ORDEM_FIXA[0]
        atividade_b = ATIVIDADES_ORDEM_FIXA[6]
        payload = self._base_payload()
        payload["atividades_selecionadas"] = [atividade_b, atividade_a]
        payload["atividades_json"] = json.dumps(
            [{"descricao": atividade_b}, {"descricao": atividade_a}],
            ensure_ascii=False,
        )
        response = self.client.post(reverse("plano_trabalho_editar", args=[self.oficio.id]), payload)
        self.assertEqual(response.status_code, 302)

        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        atividades_salvas = [item.descricao for item in plano.atividades.all().order_by("ordem", "id")]
        metas_salvas = [item.descricao for item in plano.metas.all().order_by("ordem", "id")]
        self.assertEqual(atividades_salvas, [atividade_a, atividade_b])
        self.assertEqual(
            metas_salvas,
            [META_POR_ATIVIDADE[atividade_a], META_POR_ATIVIDADE[atividade_b]],
        )

        response_get = self.client.get(
            reverse("plano_trabalho_editar", args=[self.oficio.id]),
            follow=False,
        )
        self.assertEqual(response_get.status_code, 302)
        self.assertIn(
            reverse("plano_trabalho_step1", args=[self.oficio.id]),
            response_get.url,
        )

    def test_docx_real_substitui_placeholders_paginas_01_e_02(self) -> None:
        atividade_a = ATIVIDADES_ORDEM_FIXA[0]
        atividade_b = ATIVIDADES_ORDEM_FIXA[6]
        plano = PlanoTrabalho.objects.create(
            oficio=self.oficio,
            numero=9,
            ano=2026,
            sigla_unidade="ASCOM",
            destino="Ponta Grossa/PR",
            solicitante="Joao da Silva",
            local="Praça Central",
            data_inicio=date(2026, 2, 13),
            data_fim=date(2026, 2, 16),
            horario_atendimento="09h às 17h",
            efetivo_formatado="5 servidores.",
            quantidade_servidores=5,
            composicao_diarias="4 x 100% + 1 x 30%",
            valor_unitario="150.00",
            valor_total_calculado="3225.00",
            coordenador_plano=self.coordenador,
        )
        PlanoTrabalhoAtividade.objects.create(plano=plano, ordem=1, descricao=atividade_a)
        PlanoTrabalhoAtividade.objects.create(plano=plano, ordem=2, descricao=atividade_b)
        PlanoTrabalhoMeta.objects.create(
            plano=plano,
            ordem=1,
            descricao=META_POR_ATIVIDADE[atividade_a],
        )
        PlanoTrabalhoMeta.objects.create(
            plano=plano,
            ordem=2,
            descricao=META_POR_ATIVIDADE[atividade_b],
        )
        PlanoTrabalhoRecurso.objects.create(plano=plano, ordem=1, descricao="Unidade movel da PCPR.")
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=1,
            local="Praça Central",
        )
        PlanoTrabalhoLocalAtuacao.objects.create(
            plano=plano,
            ordem=2,
            local="Ginasio Municipal",
        )

        payload = build_plano_trabalho_docx_bytes(self.oficio).getvalue()
        text = self._doc_text(payload)
        self.assertNotIn("{{dias_evento_extenso}}", text)
        self.assertNotIn("{{locais_formatado}}", text)
        self.assertNotIn("{{horario_atendimento}}", text)
        self.assertNotIn("{{quantidade_de_servidores}}", text)
        self.assertNotIn("{{solicitante}}", text)
        self.assertNotIn("{{atividades_formatada}}", text)
        self.assertNotIn("{{metas_formatada}}", text)
        self.assertIn("de 13 a 16 de fevereiro de 2026", text)
        self.assertIn("Praça Central e Ginasio Municipal", text)
        self.assertIn("das 09h às 17h", text)
        self.assertIn("Joao da Silva", text)
        self.assertIn("• Confecção da Carteira de Identidade Nacional (CIN)", text)
