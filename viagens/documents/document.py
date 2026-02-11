# viagens/documents/document.py
from __future__ import annotations

import logging
import os
import re
import shutil
import zipfile
import tempfile
from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

try:
    import pythoncom
    import win32com.client as win32client  # type: ignore
except ImportError:  # pragma: no cover - optional Windows dependency
    pythoncom = None  # type: ignore[assignment]
    win32client = None  # type: ignore[assignment]

from django.conf import settings
from django.utils import timezone
from docx import Document as DocxFactory
from docx.document import Document as DocxDocument
from num2words import num2words

from viagens.models import Cidade, ConfiguracaoOficio, Estado, Oficio, OficioConfig, Trecho, Viajante
from viagens.services.oficio_config import get_oficio_config
from viagens.services.text import title_case_pt
from viagens.utils.normalize import format_protocolo_num, format_rg


logger = logging.getLogger(__name__)

# Aceita placeholders com espaços irregulares, inclusive casos como "{{chave} }"
PLACEHOLDER_RE = re.compile(r"{{\s*([^{}]+?)\s*}(?:\s*})")

# “largura” aproximada para estimar quebra em tabela
NAME_MAX_VISUAL = 35.0
CARGO_MAX_VISUAL = 35.0

NBSP = "\u00A0"  # não quebra e “segura” linha vazia no Word

SPACE_BEFORE_PUNCT_RE = re.compile(r"\s+([,.;:!?])")
OFICIO_NUMERO_RE = re.compile(r"\bN\s*\.\s*º")
OFICIO_NUMERO_NO_SPACE_AFTER_RE = re.compile(r"\bN\.º(?=\S)")
OFICIO_NUMERO_GRAU_RE = re.compile(r"\bN\s*°")
PERCENT_DIARIAS_RE = re.compile(r"%\s*(di[aá]rias)", flags=re.IGNORECASE)
EMAIL_LABEL_RE = re.compile(r"\be-?mail:\s*", flags=re.IGNORECASE)
VIATURA_PORTE_RE = re.compile(
    r"(Viatura\s+[^\n]+?)\s+(Porte/Tr[aâ]nsito de arma:)",
    flags=re.IGNORECASE,
)

HEADER_UNIDADE_DEFAULT = "SECRETARIA DE ESTADO DA SEGURANÇA PÚBLICA"
HEADER_ORIGEM_DEFAULT = "POLÍCIA CIVIL DO PARANÁ ASSESSORIA DE COMUNICAÇÃO SOCIAL"
FOOTER_ASSINANTE_NOME_DEFAULT = "DR. RIAD BRAGA FARHAT"
FOOTER_ASSINANTE_CARGO_DEFAULT = "MD. DELEGADO GERAL ADJUNTO"
CUSTOS_SECTION_TITLE = (
    "Custos: Informar qual entidade custeara as diarias (hospedagem/alimentacao) e deslocamento:"
)


class AssinaturaObrigatoriaError(ValueError):
    pass


class DocxPdfConversionError(RuntimeError):
    pass


class MotoristaCaronaValidationError(ValueError):
    pass

FOOTER_ENDERECO_DEFAULT = (
    "Assessoria de Comunicação Social - Avenida Iguaçú, 470- Rebouças – "
    "Curitiba-PR – CEP 80.230-020:  Fone: 41-3235-6476 – e-mail:comunicacaopc.pr.gov.br"
)


# =========================
# FORMATADORES
# =========================
def _fmt_date(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%d/%m/%Y")
    except Exception:
        return str(value)


def _fmt_time(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%H:%M")
    except Exception:
        return str(value)


def _fmt_local(cidade: Cidade | None, estado: Estado | None) -> str:
    if cidade and estado:
        return f"{cidade.nome}/{estado.sigla}"
    if cidade:
        return cidade.nome
    if estado:
        return estado.sigla
    return ""


def _join(parts: Iterable[str], sep=" - ") -> str:
    return sep.join([p for p in parts if p])


def _title_case(text: str) -> str:
    return " ".join(w.capitalize() for w in (text or "").split())


def _clean_inline_text(value: str | None) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", " ").replace("\n", " ")
    return " ".join(text.split()).strip()


def _sanitize_xml_text(value: str | None) -> str:
    if value is None:
        return ""
    text = str(value)
    cleaned_chars: list[str] = []
    for ch in text:
        codepoint = ord(ch)
        if ch in ("\t", "\n", "\r"):
            cleaned_chars.append(ch)
            continue
        if (
            0x20 <= codepoint <= 0xD7FF
            or 0xE000 <= codepoint <= 0xFFFD
            or 0x10000 <= codepoint <= 0x10FFFF
        ):
            cleaned_chars.append(ch)
    return "".join(cleaned_chars)


def _sanitize_mapping_values(mapping: dict[str, str]) -> dict[str, str]:
    return {key: _sanitize_xml_text(str(value)) for key, value in mapping.items()}


def get_missing_assinatura_fields(oficio_cfg: OficioConfig) -> list[str]:
    missing: list[str] = []
    assinante = getattr(oficio_cfg, "assinante", None)
    if not assinante:
        return ["assinante"]

    if not _clean_inline_text(getattr(assinante, "nome", "") or ""):
        missing.append("nome do assinante")
    if not _clean_inline_text(getattr(assinante, "cargo", "") or ""):
        missing.append("cargo/função do assinante")
    return missing


def ensure_assinatura_config_valida(oficio_cfg: OficioConfig) -> None:
    missing = get_missing_assinatura_fields(oficio_cfg)
    if not missing:
        return
    missing_str = ", ".join(missing)
    raise AssinaturaObrigatoriaError(
        f"Não foi possível gerar o ofício. Configure a assinatura obrigatória e preencha: {missing_str}."
    )


def ensure_motorista_carona_campos(oficio: Oficio) -> None:
    if not getattr(oficio, "motorista_carona", False):
        return

    oficio_motorista = _clean_inline_text(
        getattr(oficio, "motorista_oficio_formatado", "")
        or getattr(oficio, "motorista_oficio", "")
        or ""
    )
    protocolo_motorista = _clean_inline_text(
        getattr(oficio, "motorista_protocolo_formatado", "")
        or format_protocolo_num(getattr(oficio, "motorista_protocolo", ""))
        or getattr(oficio, "motorista_protocolo", "")
        or ""
    )

    if not oficio_motorista:
        raise MotoristaCaronaValidationError("Informe o Ofício do motorista (carona).")
    if not protocolo_motorista:
        raise MotoristaCaronaValidationError("Informe o Protocolo do motorista (carona).")


def _normalize_microformat_line(text: str) -> str:
    line = _clean_inline_text(text)
    if not line:
        return ""

    line = OFICIO_NUMERO_RE.sub("N.º", line)
    line = OFICIO_NUMERO_NO_SPACE_AFTER_RE.sub("N.º ", line)
    line = OFICIO_NUMERO_GRAU_RE.sub("N.º", line)
    line = SPACE_BEFORE_PUNCT_RE.sub(r"\1", line)
    line = EMAIL_LABEL_RE.sub("e-mail: ", line)
    line = PERCENT_DIARIAS_RE.sub(r"% \1", line)
    line = re.sub(r"[ \t]{2,}", " ", line).strip()
    return line


def _normalize_microformat_text(text: str) -> str:
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = VIATURA_PORTE_RE.sub(r"\1\n\2", normalized)
    lines = [_normalize_microformat_line(part) for part in normalized.split("\n")]
    return "\n".join(lines)


def _format_diarias_quantidade(value: str | None) -> str:
    text = _normalize_microformat_line(value or "")
    if not text:
        return ""
    # Template concatena diretamente com "diárias"; garantimos um espaço antes.
    return f"{text} "


def _build_endereco_formatado(cfg: OficioConfig) -> str:
    logradouro = (getattr(cfg, "logradouro", "") or "").strip()
    numero = (getattr(cfg, "numero", "") or "").strip()
    complemento = (getattr(cfg, "complemento", "") or "").strip()
    bairro = (getattr(cfg, "bairro", "") or "").strip()
    cidade = (getattr(cfg, "cidade", "") or "").strip()
    uf = (getattr(cfg, "uf", "") or "").strip()
    cep = (getattr(cfg, "cep", "") or "").strip()

    numero_parte = ""
    if numero:
        numero_parte = numero
        if complemento:
            numero_parte = f"{numero_parte} {complemento}"

    primeira_parte = ", ".join([p for p in [logradouro, numero_parte] if p])
    cidade_uf = cidade
    if uf:
        cidade_uf = f"{cidade}/{uf}" if cidade else uf

    endereco = " - ".join([p for p in [primeira_parte, bairro, cidade_uf] if p])
    if cep:
        endereco = f"{endereco} - CEP {cep}" if endereco else f"CEP {cep}"
    return endereco


def _parse_decimal_string(value: str | Decimal | None) -> Decimal | None:
    if isinstance(value, Decimal):
        return value
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.replace("R$", "").replace(" ", "")
    text = text.replace(".", "").replace(",", ".")
    if not text:
        return None
    try:
        return Decimal(text)
    except InvalidOperation:
        return None


def _format_num(value: int, singular: str, plural: str) -> str:
    words = num2words(value, lang="pt_BR")
    unit = singular if value == 1 else plural
    return f"{words} {unit}"


def valor_por_extenso_ptbr(value: str | Decimal | None) -> str | None:
    parsed = _parse_decimal_string(value)
    if parsed is None:
        return None
    parsed = parsed.quantize(Decimal("0.01"))
    reais = int(parsed)
    centavos = int((parsed - reais) * 100)
    parts: list[str] = []
    if reais:
        parts.append(_format_num(reais, "real", "reais"))
    if centavos:
        parts.append(_format_num(centavos, "centavo", "centavos"))
    if not parts:
        parts.append("zero reais")
    return " e ".join(parts)


def is_viagem_fora_pr(trechos: list[Trecho]) -> bool:
    for trecho in trechos:
        sigla = (trecho.destino_estado.sigla if trecho.destino_estado else "") or ""
        if sigla.upper() != "PR":
            return True
    return False


def get_assunto(oficio: Oficio, trechos: list[Trecho]) -> tuple[str, str]:
    assunto_tipo = (oficio.assunto_tipo or "").strip().upper()
    if assunto_tipo == Oficio.AssuntoTipo.CONVALIDACAO:
        return (
            "Solicita\u00e7\u00e3o de convalida\u00e7\u00e3o e concess\u00e3o de di\u00e1rias.",
            "(convalida\u00e7\u00e3o)",
        )
    return ("Solicita\u00e7\u00e3o de autoriza\u00e7\u00e3o e concess\u00e3o de di\u00e1rias.", "(autoriza\u00e7\u00e3o)")


def _footer_default_parts() -> tuple[str, str, str, str]:
    text = FOOTER_ENDERECO_DEFAULT
    unidade = ""
    endereco = text
    telefone = ""
    email = ""
    if " - " in text:
        unidade, endereco = text.split(" - ", 1)
        unidade = unidade.strip()
    if "Fone:" in endereco:
        before, after = endereco.split("Fone:", 1)
        endereco = before.strip(" -: ")
        if "e-mail:" in after:
            phone_part, email_part = after.split("e-mail:", 1)
            telefone = phone_part.strip(" -: ")
            email = email_part.strip(" -: ")
        else:
            telefone = after.strip(" -: ")
    for token in ("\u2013", "\u2014"):
        endereco = endereco.replace(token, "").strip()
        telefone = telefone.replace(token, "").strip()
        email = email.replace(token, "").strip()
    return unidade, endereco.strip(), telefone.strip(), email.strip()


def _iter_docx_xml_parts_from_path(path: str) -> list[tuple[str, str]]:
    try:
        with zipfile.ZipFile(path, "r") as z:
            parts = []
            for name in z.namelist():
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                xml = z.read(name).decode("utf-8", errors="ignore")
                parts.append((name, xml))
            return parts
    except Exception:
        return []


def _iter_docx_xml_parts_from_bytes(docx_bytes: bytes) -> list[tuple[str, str]]:
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes), "r") as z:
            parts = []
            for name in z.namelist():
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                xml = z.read(name).decode("utf-8", errors="ignore")
                parts.append((name, xml))
            return parts
    except Exception:
        return []


def _extract_placeholders(parts: list[tuple[str, str]]) -> set[str]:
    keys: set[str] = set()
    for _, xml in parts:
        if "{{" not in xml:
            continue
        for match in PLACEHOLDER_RE.finditer(xml):
            keys.add(match.group(1).strip())
    return keys


def _find_unresolved_placeholders(docx_bytes: bytes) -> set[str]:
    parts = _iter_docx_xml_parts_from_bytes(docx_bytes)
    leftovers: set[str] = set()
    for _, xml in parts:
        if "{{" not in xml:
            continue
        for match in PLACEHOLDER_RE.finditer(xml):
            key = match.group(1).strip()
            if key:
                leftovers.add(key)
    return leftovers


def build_destinos_e_roteiros(oficio: Oficio, trechos: list[Trecho]) -> tuple[str, str, str]:
    sede = _fmt_local(oficio.cidade_sede, oficio.estado_sede)
    if not sede and trechos:
        sede = _fmt_local(trechos[0].origem_cidade, trechos[0].origem_estado)

    seen: set[str] = set()
    destinos: list[str] = []
    for trecho in trechos:
        destino = _fmt_local(trecho.destino_cidade, trecho.destino_estado)
        if not destino or destino == sede:
            continue
        if destino in seen:
            continue
        seen.add(destino)
        destinos.append(destino)

    destinos_str = ", ".join(destinos)
    if not destinos or not sede:
        return destinos_str, "", ""
    ida = " > ".join([sede] + destinos)
    volta = f"{destinos[-1]} > {sede}"
    return destinos_str, ida, volta


def format_tipo_viatura(tipo: str | None) -> str:
    if not tipo:
        return "-"
    normalized = tipo.strip().upper()
    if normalized == "CARACTERIZADA":
        return "Caracterizada"
    if normalized == "DESCARACTERIZADA":
        return "Descaracterizada"
    return tipo.capitalize()


def format_armamento(value: str | bool | None) -> str:
    if isinstance(value, bool):
        return "Sim" if value else "Não"
    if value is None:
        return "-"
    text = str(value).strip()
    if not text:
        return "-"
    normalized = text.upper()
    truthy = {"SIM", "S", "TRUE", "1"}
    falsy = {"NAO", "NÃO", "N", "FALSE", "0"}
    if normalized in truthy:
        return "Sim"
    if normalized in falsy:
        return "Não"
    return "Sim"


def _normalize_name(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(value.strip().split()).casefold()


def _format_carona_ref(oficio: Oficio) -> str:
    return ""


def motorista_para_documento(oficio: Oficio) -> str:
    motorista_oficio_obj = oficio.motorista_viajante
    motorista_oficio_nome = (motorista_oficio_obj.nome if motorista_oficio_obj else "") or ""
    motorista_usado_nome = _clean_inline_text((oficio.motorista or "").strip() or motorista_oficio_nome)

    if getattr(oficio, "motorista_carona", False):
        base = f"{_title_case(motorista_usado_nome)} (carona)"
        referencia = _format_carona_ref(oficio)
        if referencia:
            base = f"{base} ({referencia})"

        oficio_motorista = _clean_inline_text(
            getattr(oficio, "motorista_oficio_formatado", "")
            or getattr(oficio, "motorista_oficio", "")
            or ""
        )
        protocolo_motorista = _clean_inline_text(
            getattr(oficio, "motorista_protocolo_formatado", "")
            or format_protocolo_num(getattr(oficio, "motorista_protocolo", ""))
            or getattr(oficio, "motorista_protocolo", "")
            or ""
        )
        return "\n".join(
            [
                base,
                f"Ofício do motorista: {oficio_motorista}",
                f"Protocolo do motorista: {protocolo_motorista}",
            ]
        )

    if not motorista_oficio_nome or not motorista_usado_nome:
        same = True
    elif motorista_oficio_obj:
        same = _normalize_name(motorista_oficio_obj.nome) == _normalize_name(motorista_usado_nome)
    else:
        same = _normalize_name(motorista_oficio_nome) == _normalize_name(motorista_usado_nome)

    if same:
        return _title_case(motorista_oficio_nome or motorista_usado_nome) or "-"
    return _title_case(motorista_usado_nome) or "-"


def build_col_solicitacao(viajantes: list[Viajante], assunto_text: str) -> str:
    lines: list[str] = []
    texto = assunto_text or "-"
    for idx in range(len(viajantes)):
        lines.append(texto)
        if idx < len(viajantes) - 1:
            lines.extend(_blank_lines(1))
    return "\n".join(lines or [NBSP])


def build_col_retorno(oficio: Oficio, trechos: list[Trecho]) -> tuple[str, str]:
    saida_cidade = oficio.retorno_saida_cidade or ""
    if not saida_cidade and trechos:
        saida_cidade = _fmt_local(trechos[-1].destino_cidade, trechos[-1].destino_estado)
    chegada_cidade = oficio.retorno_chegada_cidade or ""
    if not chegada_cidade and trechos:
        chegada_cidade = _fmt_local(trechos[0].origem_cidade, trechos[0].origem_estado)

    saida_text = _join(
        [
            f"Saída {saida_cidade}:".strip(),
            _join([_fmt_date(oficio.retorno_saida_data), _fmt_time(oficio.retorno_saida_hora)], " "),
        ],
        " ",
    ).strip()
    chegada_text = _join(
        [
            f"Chegada {chegada_cidade}:".strip(),
            _join([_fmt_date(oficio.retorno_chegada_data), _fmt_time(oficio.retorno_chegada_hora)], " "),
        ],
        " ",
    ).strip()

    saida_line = saida_text or NBSP
    chegada_line = chegada_text or NBSP
    return (saida_line, chegada_line)


def get_config_oficio() -> dict[str, str]:
    defaults = ConfiguracaoOficio._default_values()
    try:
        config = ConfiguracaoOficio.get_solo()
    except Exception:
        return defaults
    return {
        "nome_chefia": config.nome_chefia or defaults["nome_chefia"],
        "cargo_chefia": config.cargo_chefia or defaults["cargo_chefia"],
        "orgao_origem": config.orgao_origem or defaults["orgao_origem"],
        "orgao_destino_padrao": config.orgao_destino_padrao
        or defaults["orgao_destino_padrao"],
    }


def _remove_paragraph(paragraph) -> None:
    # Desativado por seguranca: nao remover nos/estrutura OOXML.
    return


# Helper para substituir texto preservando o parágrafo (python-docx)
def _replace_paragraph_text(paragraph, new_text: str) -> None:
    safe_text = _sanitize_xml_text(new_text or "")
    if not paragraph.runs:
        paragraph.add_run(safe_text)
        return
    paragraph.runs[0].text = safe_text
    for run in paragraph.runs[1:]:
        run.text = ""


# =========================
# ITERADORES (parágrafos em tabelas + header/footer)
# =========================
def _iter_all_paragraphs(doc: DocxDocument):
    # body
    for p in doc.paragraphs:
        yield p
    # tables (body)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

    # headers/footers tamb?m (pra n?o perder nada)
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def _apply_config_header_footer(doc: DocxDocument, cfg: OficioConfig) -> None:
    unidade = (cfg.unidade_nome or "").strip().upper()
    origem = (cfg.origem_nome or "").strip().upper()
    rodape_unidade_title = title_case_pt(unidade)
    endereco_formatado = _build_endereco_formatado(cfg)
    telefone = (cfg.telefone or "").strip()
    email = (cfg.email or "").strip()

    rodape_line = ""
    if rodape_unidade_title or endereco_formatado or telefone or email:
        partes = [rodape_unidade_title, endereco_formatado]
        contato = []
        if telefone:
            contato.append(f"Fone: {telefone}")
        if email:
            contato.append(f"e-mail: {email}")
        if contato:
            partes.append(" - ".join(contato))
        rodape_line = " - ".join([parte for parte in partes if parte])

    for section in doc.sections:
        for p in section.header.paragraphs:
            texto = "".join(r.text for r in p.runs).strip()
            if unidade and texto == HEADER_UNIDADE_DEFAULT:
                _replace_paragraph_text(p, unidade)
            elif origem and texto == HEADER_ORIGEM_DEFAULT:
                _replace_paragraph_text(p, origem)

        for p in section.footer.paragraphs:
            texto = "".join(r.text for r in p.runs).strip()
            if rodape_line and (
                texto == FOOTER_ENDERECO_DEFAULT
                or "Assessoria de Comunica\u00e7\u00e3o Social" in texto
            ):
                _replace_paragraph_text(p, rodape_line)


def _sub_placeholders(text: str, mapping: dict[str, str]) -> str:
    if "{{" not in text:
        return _sanitize_xml_text(text)

    def repl(m):
        key = m.group(1).strip()
        # se não existir -> vazio (como você pediu)
        return _sanitize_xml_text(str(mapping.get(key, "")))

    return _sanitize_xml_text(PLACEHOLDER_RE.sub(repl, text))


def _iter_paragraphs_from_container(container):
    for p in getattr(container, "paragraphs", []):
        yield p
    for table in getattr(container, "tables", []):
        yield from _iter_paragraphs_from_table(table)


def _iter_paragraphs_from_table(table):
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_paragraphs_from_container(cell)


def replace_placeholders_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    full_text = "".join(r.text for r in paragraph.runs)
    if "{{" not in full_text:
        return

    # 1) Prefer: replace run-a-run to preserve formatting.
    changed = False
    for r in paragraph.runs:
        new_text = _sub_placeholders(r.text, mapping)
        if new_text != r.text:
            r.text = new_text
            changed = True

    if not changed:
        # try a more precise cross-run replacement
        _replace_placeholders_across_runs(paragraph, mapping)
        return

    # 2) After run-level replacement, check if anything still broken across runs
    full_text = "".join(r.text for r in paragraph.runs)
    if "{{" in full_text:
        _replace_placeholders_across_runs(paragraph, mapping)


def replace_only_placeholders(doc: DocxDocument, mapping: dict[str, str]):
    """
    Substitui somente parágrafos que contenham placeholders.
    Não toca em textos fixos (sem "{{").
    """
    safe_replace_placeholders(doc, mapping)


def safe_replace_placeholders(doc: DocxDocument, mapping: dict[str, str]) -> None:
    """
    Replace seguro de placeholders:
    - percorre parágrafos e células (body/header/footer)
    - altera somente texto dos runs
    - não remove parágrafos/tabelas/nós
    """
    safe_mapping = _sanitize_mapping_values(mapping)
    for paragraph in _iter_all_paragraphs(doc):
        replace_placeholders_in_paragraph(paragraph, safe_mapping)


def _replace_placeholders_across_runs(paragraph, mapping: dict[str, str]) -> None:
    full_text = "".join(r.text for r in paragraph.runs)
    if "{{" not in full_text:
        return

    spans: list[tuple[int, int, str]] = []
    for match in PLACEHOLDER_RE.finditer(full_text):
        key = match.group(1).strip()
        value = _sanitize_xml_text(str(mapping.get(key, "")))
        spans.append((match.start(), match.end(), value))

    if not spans:
        return

    run_bounds: list[tuple[int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        start = cursor
        cursor += len(run.text)
        run_bounds.append((start, cursor))

    # process from end to start to avoid index shifts
    for start, end, value in reversed(spans):
        first_idx = None
        last_idx = None
        for i, (rs, re) in enumerate(run_bounds):
            if rs <= start < re:
                first_idx = i
            if rs < end <= re:
                last_idx = i
                break
        if first_idx is None or last_idx is None:
            continue

        if first_idx == last_idx:
            run = paragraph.runs[first_idx]
            rs, _ = run_bounds[first_idx]
            left = run.text[: start - rs]
            right = run.text[end - rs :]
            run.text = _sanitize_xml_text(f"{left}{value}{right}")
            continue

        first_run = paragraph.runs[first_idx]
        rs_first, _ = run_bounds[first_idx]
        prefix = first_run.text[: start - rs_first]
        first_run.text = _sanitize_xml_text(f"{prefix}{value}")

        for i in range(first_idx + 1, last_idx):
            paragraph.runs[i].text = ""

        last_run = paragraph.runs[last_idx]
        rs_last, _ = run_bounds[last_idx]
        suffix = last_run.text[end - rs_last :]
        last_run.text = _sanitize_xml_text(suffix)


def _visual_length(text: str) -> float:
    total = 0.0
    for ch in text:
        total += 1.5 if ch.isupper() else 1.0
    return total


def _estimate_wrapped_lines(text: str, max_visual: float) -> int:
    words = (text or "").split()
    if not words:
        return 1

    lines = 1
    current = 0.0
    for w in words:
        wlen = _visual_length(w)
        add = wlen if current == 0 else wlen + 1.0  # espaço
        if current + add <= max_visual:
            current += add
        else:
            lines += 1
            current = wlen
    return lines


def _blank_lines(n: int) -> list[str]:
    return [NBSP for _ in range(max(0, n))]


# =========================
# BLOCOS: SERVIDORES
# =========================
def build_col_nomes(viajantes: list[Viajante]) -> str:
    lines: list[str] = []
    for i, v in enumerate(viajantes):
        nome = _title_case(_clean_inline_text(v.nome or ""))
        lines.append(nome or NBSP)

        if i < len(viajantes) - 1:
            used = _estimate_wrapped_lines(nome, NAME_MAX_VISUAL)
            blanks = 2 if used <= 1 else 1
            lines.extend(_blank_lines(blanks))

    return "\n".join(lines or [NBSP])


def build_col_rgcpf(viajantes: list[Viajante]) -> str:
    lines: list[str] = []
    for i, v in enumerate(viajantes):
        rg = _clean_inline_text(format_rg(v.rg or "")) or "-"
        cpf = _clean_inline_text(v.cpf or "") or "-"
        lines.append(f"RG: {rg}")
        lines.append(f"CPF: {cpf}")

        if i < len(viajantes) - 1:
            lines.extend(_blank_lines(1))

    return "\n".join(lines or [NBSP])


def build_col_cargo(viajantes: list[Viajante]) -> str:
    lines: list[str] = []
    for i, v in enumerate(viajantes):
        cargo = _clean_inline_text(v.cargo or "")
        lines.append(cargo or NBSP)

        if i < len(viajantes) - 1:
            used = _estimate_wrapped_lines(cargo, CARGO_MAX_VISUAL)
            blanks = 2 if used <= 1 else 1
            lines.extend(_blank_lines(blanks))

    return "\n".join(lines or [NBSP])


def _build_custos_block(oficio: Oficio) -> str:
    override = (getattr(oficio, "custeio_texto_override", "") or "").strip()
    if override:
        return f"{CUSTOS_SECTION_TITLE}\n{override}"

    options = [
        Oficio.CusteioTipoChoices.UNIDADE,
        Oficio.CusteioTipoChoices.OUTRA_INSTITUICAO,
        Oficio.CusteioTipoChoices.ONUS_LIMITADOS,
    ]
    selected = (
        (getattr(oficio, "custeio_tipo", "") or "").strip()
        or (getattr(oficio, "custos", "") or "").strip()
    )
    if not selected:
        return ""
    if selected == "SEM_ONUS":
        selected = Oficio.CusteioTipoChoices.ONUS_LIMITADOS.value
    if selected not in {opt.value for opt in options}:
        selected = Oficio.CusteioTipoChoices.UNIDADE.value

    instituicao = (getattr(oficio, "nome_instituicao_custeio", "") or "").strip()
    labels = {
        Oficio.CusteioTipoChoices.UNIDADE.value: "UNIDADE \u2013 DPC (di\u00e1ria e combust\u00edvel ser\u00e3o custeados pela DPC).",
        Oficio.CusteioTipoChoices.OUTRA_INSTITUICAO.value: "OUTRA INSTITUI\u00c7\u00c3O",
        Oficio.CusteioTipoChoices.ONUS_LIMITADOS.value: "\u00d4NUS LIMITADOS AOS PR\u00d3PRIOS VENCIMENTOS",
    }

    lines = []
    for choice in options:
        marker = "( X )" if choice.value == selected else "(   )"
        label = labels.get(choice.value, choice.label)
        if (
            choice.value == Oficio.CusteioTipoChoices.OUTRA_INSTITUICAO.value
            and choice.value == selected
            and instituicao
        ):
            label = f"{label}: {instituicao}"
        lines.append(f"{marker} {label}")

    content = "\n".join(lines).strip()
    if not content:
        return ""
    return f"{CUSTOS_SECTION_TITLE}\n{content}"


def _cleanup_motorista_optional_lines(doc: DocxDocument) -> None:
    labels = (
        "ofício do motorista",
        "oficio do motorista",
        "protocolo do motorista",
    )

    for paragraph in _iter_paragraphs_from_container(doc):
        full_text = "".join(run.text for run in paragraph.runs)
        if "motorista" not in full_text.casefold():
            continue
        if not any(label in full_text.casefold() for label in labels):
            continue

        cleaned_lines: list[str] = []
        for raw_line in full_text.splitlines():
            line = _clean_inline_text(raw_line)
            if not line:
                continue

            label, sep, value = line.partition(":")
            normalized_label = label.strip().casefold()
            if sep and normalized_label in labels:
                value = _clean_inline_text(value)
                if not value:
                    continue
                line = f"{label.strip()}: {value}"

            cleaned_lines.append(line)

        _replace_paragraph_text(paragraph, "\n".join(cleaned_lines or [NBSP]))


def _remove_footer_line_if_unidade_rodape_empty(
    doc: DocxDocument,
    mapping: dict[str, str],
) -> None:
    # Desativado por seguranca: nao remover estrutura do documento.
    return


def _remove_placeholder_line_if_empty(
    doc: DocxDocument,
    mapping: dict[str, str],
    placeholder_key: str,
) -> None:
    # Desativado por seguranca: nao remover estrutura do documento.
    return


def _remove_row(row) -> None:
    # Desativado por seguranca: nao remover estrutura do documento.
    return


def _remove_custos_section_if_empty(
    doc: DocxDocument,
    mapping: dict[str, str],
) -> None:
    # Desativado por seguranca: nao remover estrutura do documento.
    return


def _normalize_document_microformatting(doc: DocxDocument) -> None:
    targets = ("N.º", "N .º", "N°", "Nº", "diárias", "e-mail:", "Viatura", "Porte/Trânsito")

    for paragraph in _iter_all_paragraphs(doc):
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            continue
        if not any(token in full_text for token in targets):
            continue
        normalized = _normalize_microformat_text(full_text)
        if normalized != full_text:
            _replace_paragraph_text(paragraph, normalized)



# =========================
# BLOCOS: DESTINOS + ROTEIRO
# =========================
def build_destinos_bloco(trechos: list[Trecho]) -> str:
    destinos: list[str] = []
    seen = set()
    for t in trechos:
        d = _fmt_local(t.destino_cidade, t.destino_estado)
        if d and d not in seen:
            seen.add(d)
            destinos.append(d)
    return ", ".join(destinos)


def build_roteiro_ida(trechos: list[Trecho]) -> tuple[str, str]:
    saida_lines: list[str] = []
    chegada_lines: list[str] = []

    for i, t in enumerate(trechos):
        origem = _fmt_local(t.origem_cidade, t.origem_estado)
        destino = _fmt_local(t.destino_cidade, t.destino_estado)

        saida_dt = _join([_fmt_date(t.saida_data), _fmt_time(t.saida_hora)], " ")
        chegada_dt = _join([_fmt_date(t.chegada_data), _fmt_time(t.chegada_hora)], " ")

        saida_lines.append(f"Saída {origem}: {saida_dt}".strip() or NBSP)
        chegada_lines.append(f"Chegada {destino}: {chegada_dt}".strip() or NBSP)

        if i < len(trechos) - 1:
            saida_lines.extend(_blank_lines(1))
            chegada_lines.extend(_blank_lines(1))

    return (
        "\n".join(saida_lines or [NBSP]),
        "\n".join(chegada_lines or [NBSP]),
    )


# =========================
# DOCX PRINCIPAL
# =========================
def build_oficio_docx_bytes(oficio: Oficio) -> BytesIO:
    template_path = str(Path(settings.BASE_DIR) / "viagens" / "documents" / "oficio_model.docx")
    doc = DocxFactory(template_path)

    viajantes = list(oficio.viajantes.all().order_by("nome"))
    trechos = list(oficio.trechos.order_by("ordem"))  # type: ignore

    config = get_config_oficio()
    oficio_cfg = get_oficio_config()
    ensure_assinatura_config_valida(oficio_cfg)
    ensure_motorista_carona_campos(oficio)
    cfg_unidade_nome = (oficio_cfg.unidade_nome or "").strip().upper()
    cfg_origem_nome = (oficio_cfg.origem_nome or "").strip().upper()
    cfg_rodape_unidade_title = title_case_pt(cfg_unidade_nome)
    cfg_endereco_formatado = _build_endereco_formatado(oficio_cfg)
    cfg_telefone = (oficio_cfg.telefone or "").strip()
    cfg_email = (oficio_cfg.email or "").strip()
    cfg_assinante_nome_title = title_case_pt(oficio_cfg.assinante.nome or "")
    cfg_assinante_cargo_title = title_case_pt(oficio_cfg.assinante.cargo or "")

    default_unidade, default_endereco, default_telefone, default_email = _footer_default_parts()
    rodape_unidade_value = cfg_rodape_unidade_title or ""
    endereco_value = cfg_endereco_formatado or default_endereco
    telefone_value = (
        f"Fone: {cfg_telefone}"
        if cfg_telefone
        else (f"Fone: {default_telefone}" if default_telefone else "")
    )
    email_value = (
        f"e-mail: {cfg_email}"
        if cfg_email
        else (f"e-mail: {default_email}" if default_email else "")
    )

    destinos_text, roteiro_ida_text, roteiro_retorno_text = build_destinos_e_roteiros(oficio, trechos)
    assunto_linha = (oficio.assunto or "").strip()
    assunto_tipo = (oficio.assunto_tipo or "").strip().upper()
    if not assunto_tipo:
        assunto_tipo = Oficio.AssuntoTipo.AUTORIZACAO
    if assunto_linha and "convalida" in assunto_linha.lower():
        assunto_tipo = Oficio.AssuntoTipo.CONVALIDACAO
    assunto_termo = (
        "convalidação"
        if assunto_tipo == Oficio.AssuntoTipo.CONVALIDACAO
        else "autorização"
    )
    if assunto_tipo == Oficio.AssuntoTipo.CONVALIDACAO:
        assunto_text = "Solicitação de convalidação e concessão de diárias."
        assunto_oficio_text = "(Convalidação)"
    else:
        assunto_text = "Solicitação de autorização e concessão de diárias."
        assunto_oficio_text = "(AUTORIZAÇÃO)"
    if not assunto_linha:
        assunto_linha = assunto_text
    motorista_formatado = motorista_para_documento(oficio)
    tipo_viatura_text = format_tipo_viatura(oficio.tipo_viatura)
    armamento_text = format_armamento(getattr(oficio, "armamento", None))
    orgao_destino_value = oficio.get_destino_display() or Oficio.DestinoChoices.GAB.label
    orgao_origem_value = (cfg_origem_nome or config["orgao_origem"] or "").upper()
    unidade_value = cfg_unidade_nome or cfg_origem_nome or orgao_origem_value or HEADER_UNIDADE_DEFAULT
    origem_value = cfg_origem_nome or cfg_unidade_nome or orgao_origem_value or HEADER_ORIGEM_DEFAULT
    solicitacao_lines = ""
    retorno_saida_lines, retorno_chegada_lines = build_col_retorno(oficio, trechos)

    saida_lines, chegada_lines = build_roteiro_ida(trechos)

    # sede (p/ retorno) - pega da origem do 1o trecho se existir
    sede = ""
    if trechos:
        t0 = trechos[0]
        sede = _fmt_local(t0.origem_cidade, t0.origem_estado)

    # destino principal (ultimo destino do roteiro)
    destino_principal = oficio.get_destino_display()
    if not destino_principal and trechos:
        destino_principal = _fmt_local(trechos[-1].destino_cidade, trechos[-1].destino_estado)

    valor_extenso_value = (
        (oficio.valor_diarias_extenso or "").strip()
        or valor_por_extenso_ptbr(oficio.valor_diarias)
        or "(preencher manualmente)"
    )
    caracterizada_text = "Sim" if oficio.motorista_carona else "Não"

    # campos simples
    mapping = {
        "oficio": oficio.numero_formatado or oficio.oficio or "",
        "ano": str(oficio.created_at.year) if oficio.created_at else str(timezone.localdate().year),
        "data_do_oficio": _fmt_date(oficio.created_at.date()) if oficio.created_at else _fmt_date(timezone.localdate()),
        "protocolo": oficio.protocolo_formatado or oficio.protocolo or "",
        "destino": destino_principal,
        "destinos_bloco": destinos_text,
        "assunto_linha": assunto_linha,
        "assunto_termo": assunto_termo,
        "assunto": assunto_text,
        "assunto_oficio": assunto_oficio_text,
        "orgao_destino": orgao_destino_value,
        "orgao_origem": orgao_origem_value,
        "unidade": unidade_value,
        "origem": origem_value,
        "unidade_rodape": rodape_unidade_value or "",
        "endereco": endereco_value,
        "telefone": telefone_value,
        "email": email_value,
        "custo": _build_custos_block(oficio),
        "nome_chefia": cfg_assinante_nome_title,
        "cargo_chefia": cfg_assinante_cargo_title,
        "assinante_nome": cfg_assinante_nome_title,
        "assinante_cargo": cfg_assinante_cargo_title,
        "roteiro_ida": roteiro_ida_text,
        "roteiro_retorno": roteiro_retorno_text,

        "diarias_x": _format_diarias_quantidade(oficio.quantidade_diarias or ""),
        "diaria": (oficio.valor_diarias or "").strip(),
        "valor_extenso": valor_extenso_value,

        "viatura": (oficio.modelo or "").strip(),
        "tipo_viatura": tipo_viatura_text,
        "combustivel": (oficio.combustivel or "").strip(),
        "placa": (oficio.placa or "").strip(),
        "motorista": motorista_formatado,
        "motorista_formatado": motorista_formatado,
        "motorista_oficio": _clean_inline_text(
            oficio.motorista_oficio_formatado or oficio.motorista_oficio or ""
        ),
        "motorista_protocolo": _clean_inline_text(
            oficio.motorista_protocolo_formatado or oficio.motorista_protocolo or ""
        ),

        "caracterizada": caracterizada_text,
        "armamento": armamento_text,

        "sede": sede,

        # retorno
        "data_hora_saida_destino": _join(
            [_fmt_date(oficio.retorno_saida_data), _fmt_time(oficio.retorno_saida_hora)],
            " ",
        ),
        "data_hora_chegada_sede": _join(
            [_fmt_date(oficio.retorno_chegada_data), _fmt_time(oficio.retorno_chegada_hora)],
            " ",
        ),

        "motivo": (oficio.motivo or "").strip(),

        "col_servidor": build_col_nomes(viajantes),
        "col_rgcpf": build_col_rgcpf(viajantes),
        "col_cargo": build_col_cargo(viajantes),
        "col_solicitacao": solicitacao_lines,
        "col_ida_saida": saida_lines,
        "col_ida_chegada": chegada_lines,
        "col_volta_saida": retorno_saida_lines,
        "col_volta_chegada": retorno_chegada_lines,

    }

    template_placeholders = _extract_placeholders(
        _iter_docx_xml_parts_from_path(template_path)
    )
    missing = template_placeholders - set(mapping.keys())
    if missing:
        for key in missing:
            mapping[key] = ""
        if settings.DEBUG:
            logger.debug("[oficio] placeholders sem contexto: %s", sorted(missing))

    mapping = _sanitize_mapping_values(mapping)

    # substituicao segura: apenas texto de runs/paragrafos/celulas, sem remover estrutura
    safe_replace_placeholders(doc, mapping)
    _cleanup_motorista_optional_lines(doc)
    _normalize_document_microformatting(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    if settings.DEBUG:
        leftovers = _find_unresolved_placeholders(buf.getvalue())
        if leftovers:
            raw = ", ".join(sorted(leftovers))
            message = "Placeholders nao substituidos no DOCX."
            if raw:
                message += f" Encontrados: {raw}."
            raise ValueError(message)

    return buf


# =========================
# PDF (Word/Windows)
# =========================


def _ensure_pywin32_available() -> None:
    if pythoncom is None or win32client is None:
        raise RuntimeError(
            "A conversão de DOCX para PDF exige pywin32 (pythoncom + win32com). "
            "Instale a dependência no ambiente Windows para usar essa funcionalidade."
        )


DOCX_MIN_SIZE_BYTES = 5 * 1024
DOCX_REQUIRED_ZIP_ENTRIES = ("[Content_Types].xml", "word/document.xml")


def _docx_head_hex(path: Path, length: int = 16) -> str:
    if not path.exists():
        return ""
    with path.open("rb") as file_obj:
        return file_obj.read(length).hex(" ")


def _docx_diag(path: Path) -> str:
    size = path.stat().st_size if path.exists() else -1
    return f"path={path} size={size} head16={_docx_head_hex(path)}"


def _save_debug_docx_copy(
    docx_bytes: bytes,
    *,
    oficio_id: int | None,
    reason: str,
) -> Path | None:
    if not settings.DEBUG:
        return None
    debug_dir = Path(settings.BASE_DIR) / "_debug_docx"
    debug_dir.mkdir(parents=True, exist_ok=True)
    oficio_label = str(oficio_id) if oficio_id is not None else "sem_id"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    debug_path = debug_dir / f"oficio_{oficio_label}_{timestamp}_{reason}.docx"
    with debug_path.open("wb") as file_obj:
        file_obj.write(docx_bytes)
        file_obj.flush()
        os.fsync(file_obj.fileno())
    logger.error("[oficio-pdf] DOCX salvo para diagnostico: %s", debug_path)
    return debug_path


def _validate_docx_file_for_word(path: Path) -> None:
    if not path.exists():
        raise DocxPdfConversionError(
            f"DOCX temporario nao foi criado. {_docx_diag(path)}"
        )

    size = path.stat().st_size
    if size <= DOCX_MIN_SIZE_BYTES:
        raise DocxPdfConversionError(
            f"DOCX temporario muito pequeno para conversao ({size} bytes). {_docx_diag(path)}"
        )

    with path.open("rb") as file_obj:
        signature = file_obj.read(2)
    if signature != b"PK":
        raise DocxPdfConversionError(
            f"DOCX temporario sem assinatura ZIP 'PK'. {_docx_diag(path)}"
        )

    try:
        with zipfile.ZipFile(path, "r") as archive:
            names = set(archive.namelist())
            missing_entries = [name for name in DOCX_REQUIRED_ZIP_ENTRIES if name not in names]
            bad_member = archive.testzip()
    except zipfile.BadZipFile as exc:
        raise DocxPdfConversionError(
            f"DOCX temporario nao e um ZIP valido. {_docx_diag(path)}"
        ) from exc

    if missing_entries:
        raise DocxPdfConversionError(
            f"DOCX temporario sem entradas obrigatorias: {', '.join(missing_entries)}. {_docx_diag(path)}"
        )
    if bad_member is not None:
        raise DocxPdfConversionError(
            f"DOCX temporario possui entrada ZIP corrompida: {bad_member}. {_docx_diag(path)}"
        )


def _is_word_com_error(exc: Exception) -> bool:
    if pythoncom is not None and hasattr(pythoncom, "com_error"):
        try:
            if isinstance(exc, pythoncom.com_error):
                return True
        except Exception:
            pass
    return exc.__class__.__name__ == "com_error"


def _xml_error_context(xml_bytes: bytes, line: int | None, radius: int = 1) -> str:
    text = xml_bytes.decode("utf-8", errors="replace")
    lines = text.splitlines()
    if not lines:
        return ""

    if line is None:
        start = 0
        end = min(len(lines), 3)
    else:
        start = max(0, line - 1 - radius)
        end = min(len(lines), line + radius)

    snippet: list[str] = []
    for idx in range(start, end):
        marker = ">>" if line is not None and idx + 1 == line else "  "
        snippet.append(f"{marker} L{idx + 1}: {lines[idx][:260]}")
    return "\n".join(snippet)


def _collect_ooxml_targets(xml_names: set[str]) -> list[str]:
    targets: list[str] = []
    for required in ("word/document.xml",):
        if required in xml_names:
            targets.append(required)
    targets.extend(
        sorted(
            name for name in xml_names if name.startswith("word/header") and name.endswith(".xml")
        )
    )
    targets.extend(
        sorted(
            name for name in xml_names if name.startswith("word/footer") and name.endswith(".xml")
        )
    )
    for optional in ("word/styles.xml", "word/numbering.xml"):
        if optional in xml_names:
            targets.append(optional)
    return targets


def _diagnose_docx_xml_on_open_error(
    docx_path: Path,
    *,
    oficio_id: int | None,
) -> tuple[Path | None, list[str]]:
    debug_root = Path(settings.BASE_DIR) / "_debug_docx"
    debug_root.mkdir(parents=True, exist_ok=True)
    oficio_label = str(oficio_id) if oficio_id is not None else "sem_id"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    unzip_dir = debug_root / f"oficio_{oficio_label}_{timestamp}_unzipped"
    failures: list[str] = []

    try:
        with zipfile.ZipFile(docx_path, "r") as archive:
            archive.extractall(unzip_dir)
            xml_names = {
                name
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            }
            targets = _collect_ooxml_targets(xml_names)

            for target in ("word/document.xml", "word/styles.xml", "word/numbering.xml"):
                if target not in xml_names:
                    logger.error("[oficio-pdf] XML ausente no DOCX: %s", target)

            if not any(name.startswith("word/header") for name in xml_names):
                logger.error("[oficio-pdf] Nenhum header XML encontrado no DOCX.")
            if not any(name.startswith("word/footer") for name in xml_names):
                logger.error("[oficio-pdf] Nenhum footer XML encontrado no DOCX.")

            for xml_name in targets:
                raw_xml = archive.read(xml_name)
                try:
                    ET.fromstring(raw_xml)
                except ET.ParseError as parse_exc:
                    line = None
                    col = None
                    if hasattr(parse_exc, "position") and parse_exc.position:
                        line, col = parse_exc.position
                    context = _xml_error_context(raw_xml, line, radius=1)
                    failure = (
                        f"{xml_name} | erro={parse_exc} | linha={line} | coluna={col}"
                    )
                    failures.append(failure)
                    logger.error("[oficio-pdf] XML quebrado: %s", failure)
                    if context:
                        logger.error(
                            "[oficio-pdf] Contexto XML (%s):\n%s",
                            xml_name,
                            context,
                        )
    except Exception as exc:
        logger.exception("[oficio-pdf] Falha no diagnostico XML do DOCX: %s", exc)
        return None, [f"diagnostico_xml_falhou: {exc}"]

    if failures:
        logger.error(
            "[oficio-pdf] Diagnostico XML encontrou %d arquivo(s) com erro. unzip_dir=%s",
            len(failures),
            unzip_dir,
        )
    else:
        logger.error(
            "[oficio-pdf] XMLs OK; suspeita de OOXML semantico invalido. unzip_dir=%s",
            unzip_dir,
        )
    return unzip_dir, failures


def docx_bytes_to_pdf_bytes(docx_bytes: bytes, *, oficio_id: int | None = None) -> bytes:
    """
    Converte DOCX em PDF usando Microsoft Word (fidelidade alta).
    Requer Windows + Word instalado.
    """
    _ensure_pywin32_available()
    tmp_dir = Path(tempfile.mkdtemp(prefix="cv_pdf_"))
    docx_path = tmp_dir / "in.docx"
    pdf_path = tmp_dir / "out.pdf"
    keep_tmp_dir = False
    com_initialized = False
    word = None
    doc = None

    try:
        with docx_path.open("wb") as file_obj:
            file_obj.write(docx_bytes)
            file_obj.flush()
            os.fsync(file_obj.fileno())

        try:
            _validate_docx_file_for_word(docx_path)
        except Exception as exc:
            keep_tmp_dir = bool(settings.DEBUG)
            debug_path = _save_debug_docx_copy(
                docx_bytes,
                oficio_id=oficio_id,
                reason="invalid_before_word",
            )
            message = f"DOCX invalido antes da conversao no Word. {_docx_diag(docx_path)}"
            if debug_path is not None:
                message += f" debug_docx={debug_path}"
            raise DocxPdfConversionError(message) from exc

        pythoncom.CoInitialize()
        com_initialized = True

        try:
            word = win32client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            try:
                doc = word.Documents.Open(
                    FileName=str(docx_path),
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    ConfirmConversions=False,
                    Visible=False,
                    OpenAndRepair=True,
                    NoEncodingDialog=True,
                )
            except Exception as exc:
                keep_tmp_dir = bool(settings.DEBUG)
                debug_path = _save_debug_docx_copy(
                    docx_bytes,
                    oficio_id=oficio_id,
                    reason="word_open_error",
                )
                unzip_dir = None
                xml_failures: list[str] = []
                if _is_word_com_error(exc):
                    unzip_dir, xml_failures = _diagnose_docx_xml_on_open_error(
                        docx_path,
                        oficio_id=oficio_id,
                    )

                message = f"Falha ao abrir DOCX no Word COM. {_docx_diag(docx_path)}"
                if debug_path is not None:
                    message += f" debug_docx={debug_path}"
                if unzip_dir is not None:
                    message += f" unzip_dir={unzip_dir}"
                if xml_failures:
                    message += f" xml_primeiro_erro={xml_failures[0]}"
                elif unzip_dir is not None:
                    message += " xmls_ok_suspeita_ooxml_semantico_invalido=true"
                raise DocxPdfConversionError(message) from exc

            try:
                doc.ExportAsFixedFormat(
                    OutputFileName=str(pdf_path),
                    ExportFormat=17,  # PDF
                    OpenAfterExport=False,
                    OptimizeFor=0,
                    Item=0,
                )
            except Exception as exc:
                keep_tmp_dir = bool(settings.DEBUG)
                debug_path = _save_debug_docx_copy(
                    docx_bytes,
                    oficio_id=oficio_id,
                    reason="word_export_error",
                )
                message = f"Falha na exportacao DOCX->PDF via Word COM. {_docx_diag(docx_path)}"
                if debug_path is not None:
                    message += f" debug_docx={debug_path}"
                raise DocxPdfConversionError(message) from exc
        except DocxPdfConversionError:
            raise
        except Exception as exc:
            keep_tmp_dir = bool(settings.DEBUG)
            debug_path = _save_debug_docx_copy(
                docx_bytes,
                oficio_id=oficio_id,
                reason="word_com_error",
            )
            message = f"Falha na conversao DOCX->PDF via Word COM. {_docx_diag(docx_path)}"
            if debug_path is not None:
                message += f" debug_docx={debug_path}"
            raise DocxPdfConversionError(message) from exc
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    logger.exception("[oficio-pdf] Falha ao fechar documento Word.")
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    logger.exception("[oficio-pdf] Falha ao encerrar Word.Application.")

        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            keep_tmp_dir = bool(settings.DEBUG)
            debug_path = _save_debug_docx_copy(
                docx_bytes,
                oficio_id=oficio_id,
                reason="empty_pdf",
            )
            message = f"Conversao concluiu sem PDF valido. {_docx_diag(docx_path)}"
            if debug_path is not None:
                message += f" debug_docx={debug_path}"
            raise DocxPdfConversionError(message)

        return pdf_path.read_bytes()
    finally:
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                logger.exception("[oficio-pdf] Falha em CoUninitialize.")
        if tmp_dir.exists():
            if settings.DEBUG and keep_tmp_dir:
                logger.error("[oficio-pdf] Mantendo pasta temporaria para diagnostico: %s", tmp_dir)
            else:
                shutil.rmtree(tmp_dir, ignore_errors=True)


def build_oficio_docx_and_pdf_bytes(oficio: Oficio) -> tuple[bytes, bytes]:
    """
    Retorna (docx_bytes, pdf_bytes)
    """
    docx_buf = build_oficio_docx_bytes(oficio)
    docx_bytes = docx_buf.getvalue()
    pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes, oficio_id=getattr(oficio, "id", None))
    return docx_bytes, pdf_bytes


# Alias pra não quebrar import antigo em views.py
def build_oficio_docx_and_pdf(oficio: Oficio):
    return build_oficio_docx_and_pdf_bytes(oficio)
