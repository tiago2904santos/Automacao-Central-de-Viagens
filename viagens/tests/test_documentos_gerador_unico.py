from __future__ import annotations

from datetime import date, time, timedelta
from io import BytesIO
from unittest.mock import patch

from django.test import TestCase
from django.utils import timezone
from docx import Document

from viagens.documents.document import build_oficio_docx_bytes
from viagens.documents.ordem_servico import build_ordem_servico_docx_bytes
from viagens.documents.plano_trabalho import build_plano_trabalho_docx_bytes
from viagens.models import Cidade, Estado, Oficio, OficioConfig, OrdemServico, PlanoTrabalho, Trecho, Viajante


class GeradorUnicoDocumentosTests(TestCase):
    def setUp(self) -> None:
        self.estado_pr = Estado.objects.create(sigla="PR", nome="Parana")
        self.cidade_sede = Cidade.objects.create(nome="Curitiba", estado=self.estado_pr)
        self.cidade_destino = Cidade.objects.create(nome="Maringa", estado=self.estado_pr)
        self.viajante = Viajante.objects.create(
            nome="Servidor Teste",
            rg="12345678X",
            cpf="00000000000",
            cargo="Delegado",
        )
        self.assinante = Viajante.objects.create(
            nome="Maria Assinante",
            rg="99999999X",
            cpf="11111111111",
            cargo="Delegada Adjunta",
        )

    def _mock_oficio_config(self):
        cfg = OficioConfig()
        cfg.assinante = self.assinante
        cfg.unidade_nome = "SECRETARIA DE ESTADO DA SEGURANCA PUBLICA"
        cfg.origem_nome = "POLICIA CIVIL DO PARANA ASSESSORIA DE COMUNICACAO SOCIAL"
        cfg.cep = ""
        cfg.logradouro = ""
        cfg.bairro = ""
        cfg.cidade = ""
        cfg.uf = ""
        cfg.numero = ""
        cfg.complemento = ""
        cfg.telefone = ""
        cfg.email = ""
        return patch("viagens.documents.document.get_oficio_config", return_value=cfg)

    def _build_oficio(self) -> Oficio:
        oficio = Oficio.objects.create(
            oficio="123/2026",
            protocolo="456/2026",
            assunto_tipo=Oficio.AssuntoTipo.AUTORIZACAO,
            placa="ABC1234",
            modelo="Viatura Teste",
            combustivel="Gasolina",
            motorista=self.viajante.nome,
            motivo="Cobertura de evento institucional",
            tipo_destino="INTERIOR",
            quantidade_diarias="1",
            valor_diarias="290,55",
            valor_diarias_extenso="duzentos e noventa reais e cinquenta e cinco centavos",
            retorno_saida_data=date(2026, 2, 1),
            retorno_saida_hora=time(8, 0),
            retorno_chegada_data=date(2026, 2, 2),
            retorno_chegada_hora=time(18, 0),
            estado_sede=self.estado_pr,
            cidade_sede=self.cidade_sede,
            estado_destino=self.estado_pr,
            cidade_destino=self.cidade_destino,
        )
        oficio.viajantes.add(self.viajante)
        Trecho.objects.create(
            oficio=oficio,
            ordem=1,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_sede,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino,
            saida_data=date(2026, 2, 1),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 2, 1),
            chegada_hora=time(12, 0),
        )
        return oficio

    def _doc_text(self, docx_bytes: bytes) -> str:
        doc = Document(BytesIO(docx_bytes))
        chunks: list[str] = []
        for paragraph in doc.paragraphs:
            chunks.append("".join(run.text for run in paragraph.runs))
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                chunks.append("".join(run.text for run in paragraph.runs))
            for paragraph in section.footer.paragraphs:
                chunks.append("".join(run.text for run in paragraph.runs))
        return "\n".join(chunks)

    def _container_text(self, container) -> str:
        lines: list[str] = []
        for paragraph in container.paragraphs:
            lines.append("".join(run.text for run in paragraph.runs))
        return "\n".join(lines)

    def test_generate_plano_trabalho_docx(self) -> None:
        oficio = self._build_oficio()
        PlanoTrabalho.objects.create(
            oficio=oficio,
            numero=1,
            ano=2026,
            local="Curitiba/PR",
            data_inicio=date(2026, 2, 1),
            data_fim=date(2026, 2, 2),
            efetivo_por_dia=1,
            valor_total="R$ 290,55",
            coordenador_nome="Maria Assinante",
            coordenador_cargo="Delegada Adjunta",
        )

        text = self._doc_text(build_plano_trabalho_docx_bytes(oficio).getvalue())
        self.assertIn("PLANO DE TRABALHO", text)

    def test_generate_ordem_servico_docx(self) -> None:
        oficio = self._build_oficio()
        OrdemServico.objects.create(
            oficio=oficio,
            numero=2,
            ano=2026,
            referencia="Diligências",
            determinante_nome="Maria Assinante",
            determinante_cargo="Delegada Adjunta",
            finalidade="para visita técnica",
        )

        text = self._doc_text(build_ordem_servico_docx_bytes(oficio).getvalue())
        self.assertIn("ORDEM DE SERVIÇO Nº", text)

    def test_justificativa_nao_contem_destinatario(self) -> None:
        oficio = self._build_oficio()
        oficio.justificativa_texto = "Justificativa operacional do deslocamento."
        oficio.save(update_fields=["justificativa_texto"])

        trecho = oficio.trechos.get(ordem=1)
        data_saida = timezone.localdate() + timedelta(days=5)
        trecho.saida_data = data_saida
        trecho.chegada_data = data_saida
        trecho.save(update_fields=["saida_data", "chegada_data"])

        with self._mock_oficio_config():
            docx_bytes = build_oficio_docx_bytes(oficio).getvalue()

        doc = Document(BytesIO(docx_bytes))
        self.assertGreaterEqual(len(doc.sections), 2)
        justificativa_section = doc.sections[-1]
        justificativa_scope = (
            f"{self._container_text(justificativa_section.header)}\n"
            f"{self._container_text(justificativa_section.footer)}"
        ).upper()

        self.assertNotIn("EXMO", justificativa_scope)
        self.assertNotIn("EXMO.", justificativa_scope)
        self.assertNotIn("EXMO. SR:", justificativa_scope)
        self.assertNotIn("DELEGADO GERAL ADJUNTO", justificativa_scope)
