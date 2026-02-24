from __future__ import annotations

from datetime import date, time
from io import BytesIO
from unittest.mock import patch

from django.test import TestCase
from docx import Document

from viagens.documents.document import (
    AssinaturaObrigatoriaError,
    MotoristaCaronaValidationError,
    build_oficio_docx_bytes,
    build_termo_autorizacao_payload_docx_bytes,
    build_termo_autorizacao_docx_bytes,
)
from viagens.models import Cidade, Estado, Oficio, OficioConfig, Trecho, Viajante


class OficioDocumentGenerationTests(TestCase):
    def setUp(self) -> None:
        self.estado_pr = Estado.objects.create(sigla="PR", nome="Parana")
        self.cidade_sede = Cidade.objects.create(nome="Curitiba", estado=self.estado_pr)
        self.cidade_destino = Cidade.objects.create(nome="Maringa", estado=self.estado_pr)
        self.cidade_destino_2 = Cidade.objects.create(nome="Londrina", estado=self.estado_pr)
        self.cidade_destino_3 = Cidade.objects.create(nome="Ponta Grossa", estado=self.estado_pr)
        self.viajante = Viajante.objects.create(
            nome="Servidor Teste",
            rg="12345678X",
            cpf="00000000000",
            cargo="Delegado",
        )
        self.assinante = Viajante.objects.create(
            nome="Maria Assinante",
            rg="99999999X",
            cpf="11111111111",
            cargo="Delegada Adjunta",
        )

    def _build_oficio(self, assunto_tipo: str) -> Oficio:
        oficio = Oficio.objects.create(
            oficio="123/2026",
            protocolo="456/2026",
            assunto_tipo=assunto_tipo,
            placa="ABC1234",
            modelo="Viatura Teste",
            combustivel="Gasolina",
            motorista=self.viajante.nome,
            motivo="Teste",
            tipo_destino="INTERIOR",
            quantidade_diarias="1",
            valor_diarias="290,55",
            valor_diarias_extenso="duzentos e noventa reais e cinquenta e cinco centavos",
            retorno_saida_data=date(2026, 2, 1),
            retorno_saida_hora=time(8, 0),
            retorno_chegada_data=date(2026, 2, 2),
            retorno_chegada_hora=time(18, 0),
            estado_sede=self.estado_pr,
            cidade_sede=self.cidade_sede,
            estado_destino=self.estado_pr,
            cidade_destino=self.cidade_destino,
            motorista_oficio="",
            motorista_protocolo="",
        )
        oficio.viajantes.add(self.viajante)
        Trecho.objects.create(
            oficio=oficio,
            ordem=1,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_sede,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino,
            saida_data=date(2026, 2, 1),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 2, 1),
            chegada_hora=time(12, 0),
        )
        return oficio

    def _doc_text(self, docx_bytes: bytes) -> str:
        doc = Document(BytesIO(docx_bytes))
        chunks: list[str] = []
        for paragraph in doc.paragraphs:
            chunks.append("".join(run.text for run in paragraph.runs))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        chunks.append("".join(run.text for run in paragraph.runs))
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                chunks.append("".join(run.text for run in paragraph.runs))
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            chunks.append("".join(run.text for run in paragraph.runs))
            for paragraph in section.footer.paragraphs:
                chunks.append("".join(run.text for run in paragraph.runs))
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            chunks.append("".join(run.text for run in paragraph.runs))
        return "\n".join(chunks)

    def _mock_oficio_config(
        self,
        assinante: Viajante | None,
        *,
        unidade_nome: str = "",
        origem_nome: str = "",
        telefone: str = "",
        email: str = "",
    ):
        cfg = OficioConfig()
        cfg.assinante = assinante
        cfg.unidade_nome = unidade_nome
        cfg.origem_nome = origem_nome
        cfg.cep = ""
        cfg.logradouro = ""
        cfg.bairro = ""
        cfg.cidade = ""
        cfg.uf = ""
        cfg.numero = ""
        cfg.complemento = ""
        cfg.telefone = telefone
        cfg.email = email
        return patch("viagens.documents.document.get_oficio_config", return_value=cfg)

    def test_autorizacao_usa_frase_exata_sem_caixa_alta_no_corpo(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)

        with self._mock_oficio_config(self.assinante):
            docx_bytes = build_oficio_docx_bytes(oficio).getvalue()
        text = self._doc_text(docx_bytes)

        self.assertIn(
            "Senhor Delegado, Através deste, solicito autorização e medidas para a concessão de diárias e recursos para combustível, conforme cronograma abaixo:",
            text,
        )
        self.assertNotIn("solicito AUTORIZAÇÃO e medidas", text)
        self.assertIn("(AUTORIZAÇÃO)", text)

    def test_convalidacao_formata_identificacao_com_c_maiuculo(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.CONVALIDACAO)

        with self._mock_oficio_config(self.assinante):
            docx_bytes = build_oficio_docx_bytes(oficio).getvalue()
        text = self._doc_text(docx_bytes)

        self.assertIn("(Convalidação)", text)
        self.assertNotIn("(CONVALIDAÇÃO)", text)
        self.assertIn("solicito convalidação e medidas", text)

    def test_assinatura_com_assinante_e_bloqueio_sem_assinante(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)

        with self._mock_oficio_config(self.assinante):
            com_assinante = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())
        self.assertIn("Respeitosamente,", com_assinante)
        self.assertIn("Maria Assinante", com_assinante)
        self.assertIn("Delegada Adjunta", com_assinante)
        self.assertNotIn("(assinatura)", com_assinante)

        with self._mock_oficio_config(None):
            with self.assertRaises(AssinaturaObrigatoriaError) as ctx:
                build_oficio_docx_bytes(oficio)
        self.assertIn("assinante", str(ctx.exception).lower())

    def test_custos_com_titulo_fixo_quando_bloco_existe(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.custeio_tipo = Oficio.CusteioTipoChoices.UNIDADE
        oficio.save(update_fields=["custeio_tipo"])

        with self._mock_oficio_config(self.assinante):
            text = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())

        self.assertIn(
            "Custos: Informar qual entidade custeara as diarias (hospedagem/alimentacao) e deslocamento:",
            text,
        )

    def test_custos_vazio_nao_remove_estrutura_docx(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.custeio_tipo = ""
        oficio.custos = ""
        oficio.custeio_texto_override = ""
        oficio.save(update_fields=["custeio_tipo", "custos", "custeio_texto_override"])

        with self._mock_oficio_config(self.assinante):
            doc = Document(build_oficio_docx_bytes(oficio))

        self.assertNotIn(
            "Custos: Informar qual entidade custeara as diarias (hospedagem/alimentacao) e deslocamento:",
            "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells),
        )
        self.assertEqual(len(doc.tables[6].rows), 3)
        self.assertIn("Motivo da Viagem:", doc.tables[6].rows[1].cells[0].text)

    def test_col_solicitacao_fica_vazio_sem_placeholder(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        with self._mock_oficio_config(self.assinante):
            docx_buf = build_oficio_docx_bytes(oficio)
            doc = Document(docx_buf)
            text = self._doc_text(docx_buf.getvalue())

        solicitacao_cell = doc.tables[1].rows[1].cells[3]
        self.assertEqual(solicitacao_cell.text.strip(), "")
        self.assertNotIn("{{col_solicitacao}}", text)

    def test_carona_imprime_motorista_oficio_e_protocolo(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.motorista_carona = True
        oficio.motorista = "Maria Assinante"
        oficio.motorista_oficio = "123/2026"
        oficio.motorista_protocolo = "456/2026"
        oficio.save(
            update_fields=[
                "motorista_carona",
                "motorista",
                "motorista_oficio",
                "motorista_protocolo",
            ]
        )

        with self._mock_oficio_config(self.assinante):
            text = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())

        self.assertIn("Motorista: Maria Assinante (carona)", text)
        self.assertIn("Ofício do motorista: 123/2026", text)
        self.assertIn("Protocolo do motorista: 4562026", text)

    def test_carona_sem_oficio_bloqueia_geracao(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.motorista_carona = True
        oficio.motorista = "Maria Assinante"
        oficio.motorista_oficio = ""
        oficio.motorista_protocolo = "456/2026"
        oficio.save(
            update_fields=[
                "motorista_carona",
                "motorista",
                "motorista_oficio",
                "motorista_protocolo",
            ]
        )

        with self._mock_oficio_config(self.assinante):
            with self.assertRaises(MotoristaCaronaValidationError) as ctx:
                build_oficio_docx_bytes(oficio)
        self.assertEqual(str(ctx.exception), "Informe o Ofício do motorista (carona).")

    def test_carona_sem_protocolo_bloqueia_geracao(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.motorista_carona = True
        oficio.motorista = "Maria Assinante"
        oficio.motorista_oficio = "123/2026"
        oficio.motorista_protocolo = ""
        oficio.save(
            update_fields=[
                "motorista_carona",
                "motorista",
                "motorista_oficio",
                "motorista_protocolo",
            ]
        )

        with self._mock_oficio_config(self.assinante):
            with self.assertRaises(MotoristaCaronaValidationError) as ctx:
                build_oficio_docx_bytes(oficio)
        self.assertEqual(str(ctx.exception), "Informe o Protocolo do motorista (carona).")

    def test_remove_linhas_motorista_quando_campos_estao_vazios(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.motorista_oficio = ""
        oficio.motorista_protocolo = ""
        oficio.save(update_fields=["motorista_oficio", "motorista_protocolo"])

        with self._mock_oficio_config(self.assinante):
            text = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())
        self.assertNotIn("Ofício do motorista:", text)
        self.assertNotIn("Protocolo do motorista:", text)

    def test_microformatting_oficio_diarias_email_e_viatura_porte(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        oficio.quantidade_diarias = "1 x 30%"
        oficio.tipo_viatura = "DESCARACTERIZADA"
        oficio.save(update_fields=["quantidade_diarias", "tipo_viatura"])

        with self._mock_oficio_config(
            self.assinante,
            unidade_nome="ASSESSORIA DE COMUNICACAO",
            origem_nome="POLICIA CIVIL",
            email="comunicacao@pc.pr.gov.br",
        ):
            text = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())

        self.assertIn("Ofício N.º 123/2026 (AUTORIZAÇÃO)", text)
        self.assertNotIn("Ofício N .º", text)
        self.assertIn("1 x 30% diárias", text)
        self.assertIn("e-mail: comunicacao@pc.pr.gov.br", text)
        self.assertNotIn("Viatura Descaracterizada Porte/Trânsito de arma", text)

    def test_footer_sem_placeholder_pendente_quando_unidade_rodape_vazio(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)

        with self._mock_oficio_config(self.assinante):
            text = self._doc_text(build_oficio_docx_bytes(oficio).getvalue())

        self.assertNotIn("{{unidade_rodape}", text)
        self.assertNotIn("{{unidade_rodape} }", text)
        self.assertIn("e-mail:", text)

    def test_termo_autorizacao_formata_data_extenso_e_destinos_multiplos(self) -> None:
        oficio = self._build_oficio(Oficio.AssuntoTipo.AUTORIZACAO)
        Trecho.objects.create(
            oficio=oficio,
            ordem=2,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_destino,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino_2,
            saida_data=date(2026, 2, 1),
            saida_hora=time(13, 0),
            chegada_data=date(2026, 2, 1),
            chegada_hora=time(16, 0),
        )
        Trecho.objects.create(
            oficio=oficio,
            ordem=3,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_destino_2,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_destino_3,
            saida_data=date(2026, 2, 2),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 2, 2),
            chegada_hora=time(11, 0),
        )

        with self._mock_oficio_config(
            self.assinante,
            unidade_nome="POLICIA CIVIL DO PARANA",
            origem_nome="ASCOM",
            telefone="(41) 3235-6476",
            email="ascom@pc.pr.gov.br",
        ):
            docx_bytes = build_termo_autorizacao_docx_bytes(oficio).getvalue()

        text = self._doc_text(docx_bytes)
        self.assertIn("dia 1 a 2 de fevereiro de 2026", text)
        self.assertIn("Maringa/PR, Londrina/PR e Ponta Grossa/PR", text)
        self.assertIn("ASCOM", text)
        self.assertIn("POLICIA CIVIL DO PARANA", text)
        self.assertNotIn("{{data_do_evento}}", text)
        self.assertNotIn("{{destino}}", text)

    def test_termo_payload_formata_datas_multiplas(self) -> None:
        with self._mock_oficio_config(
            self.assinante,
            unidade_nome="POLICIA CIVIL DO PARANA",
            origem_nome="ASCOM",
        ):
            docx_bytes = build_termo_autorizacao_payload_docx_bytes(
                datas=[date(2026, 2, 13), date(2026, 2, 15), date(2026, 2, 16)],
                destinos=["Curitiba/PR", "Maringa/PR"],
            ).getvalue()

        text = self._doc_text(docx_bytes)
        self.assertIn("dias 13, 15 e 16 de fevereiro de 2026", text)
        self.assertIn("Curitiba/PR e Maringa/PR", text)
