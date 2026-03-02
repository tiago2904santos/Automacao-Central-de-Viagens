from __future__ import annotations

import json
from datetime import date, time
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.test import TestCase
from django.urls import reverse
from docx import Document

from viagens.documents.plano_trabalho import build_plano_trabalho_docx_bytes
from viagens.forms import PlanoTrabalhoStep1Form
from viagens.models import Cargo, Cidade, Estado, Oficio, PlanoTrabalho, Trecho, Viajante
from viagens.services.oficio_config import get_oficio_config
from viagens.services.diarias_unified import parse_decimal_br
from viagens.services.plano_trabalho import (
    DEFAULT_COORDENADOR_PLANO_NOME,
    build_plano_placeholders,
    efetivo_total_servidores,
    format_horario_atendimento,
    format_total_servidores,
    has_coordenador_municipal,
)


class PlanoTrabalhoWizardTests(TestCase):
    def setUp(self) -> None:
        self.estado_pr = Estado.objects.create(sigla="PR", nome="Parana")
        self.cidade_curitiba = Cidade.objects.create(nome="Curitiba", estado=self.estado_pr)
        self.cidade_foz = Cidade.objects.create(nome="Foz do Iguacu", estado=self.estado_pr)
        self.cidade_londrina = Cidade.objects.create(nome="Londrina", estado=self.estado_pr)
        self.servidor = Viajante.objects.create(
            nome="Servidor Coordenador",
            rg="12345678X",
            cpf="00000000000",
            cargo="Delegado",
        )
        self.cargo_delegado = Cargo.objects.create(nome="Delegado")
        self.cargo_escrivao = Cargo.objects.create(nome="Escrivao")
        self.oficio = Oficio.objects.create(
            oficio="901/2026",
            protocolo="121234567",
            assunto="Plano de trabalho",
            placa="ABC1234",
            modelo="Viatura",
            combustivel="Gasolina",
            motorista=self.servidor.nome,
            tipo_destino="INTERIOR",
            motivo="Atendimento institucional",
            estado_sede=self.estado_pr,
            cidade_sede=self.cidade_curitiba,
            estado_destino=self.estado_pr,
            cidade_destino=self.cidade_foz,
            retorno_chegada_data=date(2026, 3, 14),
            retorno_chegada_hora=time(18, 0),
        )
        self.oficio.viajantes.add(self.servidor)
        Trecho.objects.create(
            oficio=self.oficio,
            ordem=1,
            origem_estado=self.estado_pr,
            origem_cidade=self.cidade_curitiba,
            destino_estado=self.estado_pr,
            destino_cidade=self.cidade_foz,
            saida_data=date(2026, 3, 12),
            saida_hora=time(8, 0),
            chegada_data=date(2026, 3, 12),
            chegada_hora=time(12, 0),
        )
        self.atividade_a = "Atividade institucional A"
        self.atividade_b = "Atividade institucional B"
        self.meta_a = "Meta institucional A"
        self.meta_b = "Meta institucional B"

    def _post_step1(self, *, solicitantes: list[str], nome_pcpr: str = "") -> None:
        response = self.client.post(
            reverse("plano_trabalho_step1", args=[self.oficio.id]),
            {
                "solicitantes": solicitantes,
                "solicitante_pcpr_nome": nome_pcpr,
                "data_unica": "",
                "data_inicio": "2026-03-12",
                "data_fim": "2026-03-14",
                "horario_inicio": "09:00",
                "horario_fim": "17:00",
                "destinos-TOTAL_FORMS": "2",
                "destinos-order": "0,1",
                "destinos-0-uf": "PR",
                "destinos-0-cidade": str(self.cidade_foz.id),
                "destinos-1-uf": "PR",
                "destinos-1-cidade": str(self.cidade_londrina.id),
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse("plano_trabalho_step2", args=[self.oficio.id]), response.url)

    def _doc_text(self, payload: bytes) -> str:
        doc = Document(BytesIO(payload))
        chunks = [paragraph.text for paragraph in doc.paragraphs]
        for section in doc.sections:
            chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
            chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(chunks)

    def _conteudo_plano_payload(
        self,
        *,
        atividades: list[str] | None = None,
        metas: list[str] | None = None,
    ) -> dict[str, str]:
        atividades_values = atividades if atividades is not None else [self.atividade_a]
        metas_values = metas if metas is not None else [self.meta_a]
        return {
            "atividades_json": json.dumps(
                [{"descricao": item} for item in atividades_values],
                ensure_ascii=False,
            ),
            "metas_json": json.dumps(
                [{"descricao": item} for item in metas_values],
                ensure_ascii=False,
            ),
        }

    def test_step1_persistencia_campos_evento(self) -> None:
        self._post_step1(solicitantes=["PCPR na Comunidade"], nome_pcpr="Prefeitura Municipal")
        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        self.assertEqual(plano.solicitantes_json, ["PCPR na Comunidade"])
        self.assertEqual(plano.solicitante, "Prefeitura Municipal")
        self.assertEqual(plano.data_inicio, date(2026, 3, 12))
        self.assertEqual(plano.data_fim, date(2026, 3, 14))
        self.assertEqual(plano.horario_atendimento, "das 09h às 17h")
        self.assertEqual(len(plano.destinos_json), 2)

    def test_step2_sem_coordenador_municipal_quando_nao_pcpr(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [
                        {"cargo_id": self.cargo_delegado.id, "quantidade": 1},
                        {"cargo_id": self.cargo_escrivao.id, "quantidade": 2},
                    ]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": "IGNORAR",
                "coordenador_plano_cargo": "IGNORAR",
            },
        )
        self.assertEqual(response.status_code, 302)
        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        self.assertEqual(plano.quantidade_servidores, 3)
        self.assertFalse(plano.possui_coordenador_municipal)
        self.assertIsNone(plano.coordenador_municipal)

    def test_step2_com_coordenador_municipal_quando_pcpr(self) -> None:
        self._post_step1(solicitantes=["PCPR na Comunidade"], nome_pcpr="Prefeitura")
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 1}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "sim",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "Carlos Lima",
                "coordenador_municipal_cargo": "Secretario Municipal",
                "coordenador_municipal_cidade": "Foz do Iguacu",
            },
        )
        self.assertEqual(response.status_code, 302)
        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        self.assertTrue(plano.possui_coordenador_municipal)
        self.assertIsNotNone(plano.coordenador_municipal)
        self.assertEqual(plano.coordenador_municipal.nome, "Carlos Lima")
        self.assertTrue(plano.unidade_movel)

    def test_step2_exige_coordenador_municipal_quando_pcpr(self) -> None:
        self._post_step1(solicitantes=["PCPR na Comunidade"], nome_pcpr="Prefeitura")
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 1}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "",
                "coordenador_municipal_cargo": "",
                "coordenador_municipal_cidade": "",
            },
        )
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("coordenador_municipal", form.errors)

    def test_step2_default_coordenador_administrativo_juliana(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response = self.client.get(reverse("plano_trabalho_step2", args=[self.oficio.id]))
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertEqual(form.initial["coordenador_plano_nome"], DEFAULT_COORDENADOR_PLANO_NOME)

    def test_step3_resumo_e_docx_sem_placeholders(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 2}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "sim",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "",
                "coordenador_municipal_cargo": "",
                "coordenador_municipal_cidade": "",
            },
        )
        self.assertEqual(response_step2.status_code, 302)

        response_step3 = self.client.post(
            reverse("plano_trabalho_step3", args=[self.oficio.id]),
            {
                "composicao_diarias": "2 x 100%",
                "valor_unitario": "290,55",
                "valor_total_calculado": "581,10",
                "recursos_json": json.dumps(
                    [
                        {"descricao": "Computadores"},
                        {"descricao": "Kits biometricos"},
                    ]
                ),
            },
        )
        self.assertEqual(response_step3.status_code, 302)
        self.assertIn(reverse("plano_trabalho_resumo", args=[self.oficio.id]), response_step3.url)

        resumo = self.client.get(reverse("plano_trabalho_resumo", args=[self.oficio.id]))
        self.assertEqual(resumo.status_code, 200)
        self.assertContains(resumo, "Valor por servidor")

        payload = build_plano_trabalho_docx_bytes(self.oficio).getvalue()
        text = self._doc_text(payload)
        self.assertNotIn("{{", text)
        self.assertNotIn("}}", text)
        self.assertIn("das 09h às 17h", text)

    def test_step1_aceita_horario_com_virada_de_dia(self) -> None:
        response = self.client.post(
            reverse("plano_trabalho_step1", args=[self.oficio.id]),
            {
                "solicitantes": ["PCPR na Comunidade"],
                "solicitante_pcpr_nome": "Prefeitura Municipal",
                "data_unica": "",
                "data_inicio": "2026-03-12",
                "data_fim": "2026-03-14",
                "horario_inicio": "17:00",
                "horario_fim": "02:00",
                "destinos-TOTAL_FORMS": "1",
                "destinos-order": "0",
                "destinos-0-uf": "PR",
                "destinos-0-cidade": str(self.cidade_foz.id),
            },
        )
        self.assertEqual(response.status_code, 302)
        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        self.assertEqual(plano.horario_atendimento, "das 17h às 02h do dia seguinte")

    def test_step1_form_invalido_quando_horarios_iguais(self) -> None:
        form = PlanoTrabalhoStep1Form(
            data={
                "solicitantes": ["PCPR na Comunidade"],
                "solicitante_pcpr_nome": "Prefeitura",
                "data_unica": "",
                "data_inicio": "2026-03-12",
                "data_fim": "2026-03-13",
                "horario_inicio": "17:00",
                "horario_fim": "17:00",
            }
        )
        self.assertFalse(form.is_valid())
        self.assertIn("horario_fim", form.errors)
        self.assertIn("nao pode ser igual", form.errors["horario_fim"][0])

    def test_format_horario_atendimento_com_e_sem_dia_seguinte(self) -> None:
        self.assertEqual(
            format_horario_atendimento(time(8, 0), time(17, 0)),
            "das 08h às 17h",
        )
        self.assertEqual(
            format_horario_atendimento(time(17, 0), time(2, 0)),
            "das 17h às 02h do dia seguinte",
        )
        self.assertEqual(
            format_horario_atendimento(time(22, 0), time(1, 0)),
            "das 22h às 01h do dia seguinte",
        )

    def test_format_total_servidores_pluralizacao(self) -> None:
        self.assertEqual(format_total_servidores(1), "1 servidor")
        self.assertEqual(format_total_servidores(2), "2 servidores")

    def test_regra_automatica_coordenador_municipal_por_solicitante(self) -> None:
        self.assertTrue(has_coordenador_municipal(["PCPR na Comunidade"]))
        self.assertFalse(has_coordenador_municipal(["Parana em Acao"]))
        self.assertTrue(
            has_coordenador_municipal(["Parana em Acao", "PCPR na Comunidade"])
        )

    def test_efetivo_total_soma_itens_validos(self) -> None:
        total = efetivo_total_servidores(
            [
                {"cargo_id": self.cargo_delegado.id, "quantidade": 1},
                {"cargo_id": self.cargo_escrivao.id, "quantidade": 2},
                {"cargo_id": self.cargo_delegado.id, "quantidade": 0},
            ]
        )
        self.assertEqual(total, 3)

    def test_step2_bloqueia_cargo_duplicado(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [
                        {"cargo_id": self.cargo_delegado.id, "quantidade": 1},
                        {"cargo_id": self.cargo_delegado.id, "quantidade": 2},
                    ]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": "IGNORAR",
                "coordenador_plano_cargo": "IGNORAR",
            },
        )
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("efetivo_json", form.errors)
        self.assertIn("Cargo ja adicionado", form.errors["efetivo_json"][0])

    def test_step2_persistencia_mantem_itens_ao_voltar(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [
                        {"cargo_id": self.cargo_delegado.id, "quantidade": 1},
                        {"cargo_id": self.cargo_escrivao.id, "quantidade": 2},
                    ]
                ),
                **self._conteudo_plano_payload(
                    atividades=[self.atividade_a, self.atividade_b],
                    metas=[self.meta_a, self.meta_b],
                ),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": "IGNORAR",
                "coordenador_plano_cargo": "IGNORAR",
            },
        )
        self.assertEqual(response_step2.status_code, 302)
        self.assertIn(reverse("plano_trabalho_step3", args=[self.oficio.id]), response_step2.url)

        response_back = self.client.get(reverse("plano_trabalho_step2", args=[self.oficio.id]))
        self.assertEqual(response_back.status_code, 200)
        payload_efetivo = json.loads(response_back.context["form"].initial["efetivo_json"])
        self.assertEqual(
            payload_efetivo,
            [
                {
                    "cargo_id": self.cargo_delegado.id,
                    "cargo": self.cargo_delegado.nome,
                    "quantidade": 1,
                },
                {
                    "cargo_id": self.cargo_escrivao.id,
                    "cargo": self.cargo_escrivao.nome,
                    "quantidade": 2,
                },
            ],
        )
        payload_atividades = json.loads(response_back.context["form"].initial["atividades_json"])
        payload_metas = json.loads(response_back.context["form"].initial["metas_json"])
        self.assertEqual(
            payload_atividades,
            [{"descricao": self.atividade_a}, {"descricao": self.atividade_b}],
        )
        self.assertEqual(
            payload_metas,
            [{"descricao": self.meta_a}, {"descricao": self.meta_b}],
        )

    def test_step2_bloqueia_item_vazio_em_metas_ou_atividades(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 1}]
                ),
                "atividades_json": json.dumps(
                    [{"descricao": "Atividade valida"}, {"descricao": ""}],
                    ensure_ascii=False,
                ),
                "metas_json": json.dumps(
                    [{"descricao": "Meta valida"}],
                    ensure_ascii=False,
                ),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
            },
        )
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("atividades_json", form.errors)
        self.assertIn("Remova itens vazios", form.errors["atividades_json"][0])

    def test_step2_bloqueia_quando_sem_metas(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 1}]
                ),
                "atividades_json": json.dumps(
                    [{"descricao": "Atividade valida"}],
                    ensure_ascii=False,
                ),
                "metas_json": json.dumps([], ensure_ascii=False),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
            },
        )
        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("metas_json", form.errors)
        self.assertIn("Informe ao menos 1 meta", form.errors["metas_json"][0])

    def test_step3_get_nao_quebra_quando_diarias_ainda_nao_podem_ser_calculadas(self) -> None:
        self.oficio.trechos.all().delete()

        response = self.client.get(reverse("plano_trabalho_step3", args=[self.oficio.id]))

        self.assertEqual(response.status_code, 200)
        totais = response.context["diarias_resultado"]["totais"]
        self.assertIsInstance(totais, dict)
        self.assertEqual(totais["total_geral"], "")
        self.assertEqual(totais["total_valor"], "")
        self.assertIn("valor_por_servidor", totais)
        self.assertIn("diarias_por_servidor", totais)
        self.assertContains(response, "Nao ha trechos validos para calcular as diarias do plano.")

    def test_step3_endpoint_calculo_diarias_retorna_totais_dict_mesmo_em_erro(self) -> None:
        self.oficio.trechos.all().delete()

        response = self.client.get(
            reverse("plano_trabalho_calcular_diarias", args=[self.oficio.id])
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertEqual(payload["error"], "Nao ha trechos validos para calcular as diarias do plano.")
        self.assertIsInstance(payload["totais"], dict)
        self.assertEqual(payload["totais"]["total_geral"], "")
        self.assertEqual(payload["totais"]["total_valor"], "")
        self.assertIn("valor_por_servidor", payload["totais"])
        self.assertIn("diarias_por_servidor", payload["totais"])

    def test_step3_endpoint_calculo_diarias_retorna_valor_por_servidor(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 3}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "",
                "coordenador_municipal_cargo": "",
                "coordenador_municipal_cidade": "",
            },
        )
        self.assertEqual(response_step2.status_code, 302)

        response = self.client.get(
            reverse("plano_trabalho_calcular_diarias", args=[self.oficio.id])
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        totais = payload["totais"]
        self.assertIn("valor_por_servidor", totais)
        self.assertIn("valor_unitario", totais)
        self.assertIn("total_geral", totais)
        self.assertTrue(totais["valor_por_servidor"])
        self.assertTrue(totais["total_geral"])

    def test_docx_deriva_financeiro_mesmo_sem_campos_manuais(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 2}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "sim",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "",
                "coordenador_municipal_cargo": "",
                "coordenador_municipal_cidade": "",
            },
        )
        self.assertEqual(response_step2.status_code, 302)

        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        plano.composicao_diarias = ""
        plano.valor_unitario = None
        plano.valor_total_calculado = None
        plano.save(update_fields=["composicao_diarias", "valor_unitario", "valor_total_calculado", "updated_at"])

        payload = build_plano_trabalho_docx_bytes(self.oficio).getvalue()
        self.assertTrue(payload)
        plano.refresh_from_db()
        self.assertTrue(plano.composicao_diarias)
        self.assertIsNotNone(plano.valor_total_calculado)
        self.assertIsNotNone(plano.valor_unitario)
        self.assertGreater(parse_decimal_br(str(plano.valor_total_calculado)) or 0, 0)

    def test_contexto_docx_metas_e_atividades_nao_usam_hifen(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 2}]
                ),
                **self._conteudo_plano_payload(
                    atividades=[self.atividade_a, self.atividade_b],
                    metas=[self.meta_a, self.meta_b],
                ),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
            },
        )
        self.assertEqual(response_step2.status_code, 302)
        plano = PlanoTrabalho.objects.get(oficio=self.oficio)
        placeholders = build_plano_placeholders(plano, self.oficio, get_oficio_config())
        self.assertNotEqual(placeholders["atividades_formatada"], "-")
        self.assertNotEqual(placeholders["metas_formatadas"], "-")
        self.assertIn("• Atividade institucional A", placeholders["atividades_formatada"])
        self.assertIn("• Meta institucional A", placeholders["metas_formatadas"])

    def test_rota_legada_editar_redireciona_para_step1_quando_incompleto(self) -> None:
        response = self.client.get(
            reverse("plano_trabalho_editar", args=[self.oficio.id]),
            follow=False,
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse("plano_trabalho_step1", args=[self.oficio.id]), response.url)

    def test_rota_legada_editar_redireciona_para_resumo_quando_completo(self) -> None:
        self._post_step1(solicitantes=["Parana em Acao"])
        response_step2 = self.client.post(
            reverse("plano_trabalho_step2", args=[self.oficio.id]),
            {
                "efetivo_json": json.dumps(
                    [{"cargo_id": self.cargo_delegado.id, "quantidade": 2}]
                ),
                **self._conteudo_plano_payload(),
                "unidade_movel": "nao",
                "coordenador_plano": str(self.servidor.id),
                "coordenador_plano_nome": self.servidor.nome,
                "coordenador_plano_cargo": self.servidor.cargo,
                "coordenador_municipal": "",
                "coordenador_municipal_nome": "",
                "coordenador_municipal_cargo": "",
                "coordenador_municipal_cidade": "",
            },
        )
        self.assertEqual(response_step2.status_code, 302)
        response_step3 = self.client.post(
            reverse("plano_trabalho_step3", args=[self.oficio.id]),
            {
                "composicao_diarias": "2 x 100%",
                "valor_unitario": "290,55",
                "valor_total_calculado": "581,10",
                "recursos_json": json.dumps([{"descricao": "Computadores"}]),
            },
        )
        self.assertEqual(response_step3.status_code, 302)

        response = self.client.get(
            reverse("plano_trabalho_editar", args=[self.oficio.id]),
            follow=False,
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse("plano_trabalho_resumo", args=[self.oficio.id]), response.url)

    def test_templates_wizard_sem_referencia_etapa_6(self) -> None:
        template_dir = Path(settings.BASE_DIR) / "viagens" / "templates" / "viagens"
        for name in (
            "plano_trabalho_step1.html",
            "plano_trabalho_step2.html",
            "plano_trabalho_step3.html",
            "plano_trabalho_step4.html",
        ):
            content = (template_dir / name).read_text(encoding="utf-8").lower()
            self.assertNotIn("etapa 6", content)
            self.assertNotIn("etapa-6", content)
            self.assertNotIn("revisao e geracao", content)
